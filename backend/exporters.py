#!/usr/bin/env python3
"""
Export modules for Earnings Script Analyzer
- PDF export with color-highlighted script
- Word export with track changes for suggested edits
"""

import re
from pathlib import Path
from datetime import datetime

# ============================================================
# REWRITE SUGGESTIONS
# ============================================================

REWRITE_SUGGESTIONS = {
    # Hedging phrases → confident alternatives
    'going forward': 'in the coming quarter',
    'at this time': 'currently',
    'cautiously optimistic': 'optimistic',
    'challenging environment': 'competitive market',
    'headwinds': 'market pressures',
    'kind of': '',
    'sort of': '',
    'more or less': '',
    'in some ways': '',
    'to some extent': 'partially',
    
    # False consensus markers
    'you know': '',
    'as you know': '',
    'everyone knows': '',
    'obviously': '',
    'of course': '',
    'frankly': '',
    'honestly': '',
    'to be honest': '',
    'the fact is': '',
    'the reality is': '',
    
    # Distancing → ownership
    'the company believes': 'we believe',
    'the company expects': 'we expect',
    'the company is committed': 'we are committed',
    'management believes': 'we believe',
    'management expects': 'we expect',
    'the team believes': 'we believe',
    'the team expects': 'we expect',
}

def get_suggested_rewrite(sentence):
    """Generate a suggested rewrite for a flagged sentence"""
    rewritten = sentence
    changes_made = []
    
    for pattern, replacement in REWRITE_SUGGESTIONS.items():
        if pattern.lower() in rewritten.lower():
            # Case-insensitive replacement
            regex = re.compile(re.escape(pattern), re.IGNORECASE)
            if replacement:
                rewritten = regex.sub(replacement, rewritten)
                changes_made.append(f"'{pattern}' → '{replacement}'")
            else:
                rewritten = regex.sub('', rewritten)
                changes_made.append(f"removed '{pattern}'")
    
    # Clean up extra spaces
    rewritten = re.sub(r'\s+', ' ', rewritten).strip()
    rewritten = re.sub(r'\s+([.,;:!?])', r'\1', rewritten)
    
    if changes_made:
        return rewritten, changes_made
    return None, []


# ============================================================
# SENTENCE CLASSIFICATION
# ============================================================

HEDGING_WORDS = {
    'might', 'may', 'could', 'possibly', 'perhaps', 'potentially',
    'somewhat', 'relatively', 'approximately', 'around', 'about',
    'generally', 'typically', 'usually', 'sometimes', 'occasionally',
    'likely', 'unlikely', 'probable', 'possible', 'uncertain',
    'believe', 'think', 'feel', 'hope', 'expect', 'anticipate',
    'estimate', 'appear', 'seem', 'suggest', 'indicate',
}

HEDGING_PHRASES = [
    'going forward', 'at this time', 'as you know', 'you know',
    'kind of', 'sort of', 'more or less', 'to some extent',
    'challenging environment', 'headwinds', 'cautiously optimistic',
]

DECEPTION_MARKERS = [
    'you know', 'as you know', 'everyone knows', 'obviously',
    'of course', 'clearly', 'frankly', 'honestly',
    'to be honest', 'truthfully', 'the fact is', 'the reality is',
]

DISTANCING_PHRASES = [
    'the company', 'the team', 'the organization', 
    'management', 'the business', 'the firm',
]

EXTREME_POSITIVE = {
    'tremendous', 'incredible', 'fantastic', 'amazing', 'extraordinary',
    'exceptional', 'outstanding', 'remarkable', 'phenomenal', 'spectacular',
}

CERTAINTY_WORDS = {
    'will', 'shall', 'must', 'definitely', 'certainly', 'absolutely',
    'confident', 'committed', 'guaranteed', 'assured', 'determined',
}


def classify_sentence(sentence):
    """
    Classify a sentence as RED (problem), YELLOW (watch), or GREEN (good)
    Returns (color, issues)
    """
    sentence_lower = sentence.lower()
    words = re.findall(r"\b[\w']+\b", sentence_lower)
    issues = []
    
    # Count problems
    hedge_count = sum(1 for w in words if w in HEDGING_WORDS)
    hedge_count += sum(1 for phrase in HEDGING_PHRASES if phrase in sentence_lower)
    
    deception = sum(1 for phrase in DECEPTION_MARKERS if phrase in sentence_lower)
    
    distancing = sum(1 for phrase in DISTANCING_PHRASES if phrase in sentence_lower)
    has_we = 'we ' in sentence_lower or 'our ' in sentence_lower
    
    extreme = sum(1 for w in words if w in EXTREME_POSITIVE)
    
    certainty = sum(1 for w in words if w in CERTAINTY_WORDS)
    
    # Classify
    if hedge_count >= 3:
        issues.append(f"Triple hedge ({hedge_count} markers)")
    if deception >= 1:
        issues.append("False consensus marker")
    if distancing >= 1 and not has_we:
        issues.append("Distancing language")
    if extreme >= 2:
        issues.append("Excessive superlatives")
    
    # Determine color
    if len(issues) >= 2 or hedge_count >= 4 or deception >= 2:
        return 'RED', issues
    elif len(issues) >= 1 or hedge_count >= 2:
        return 'YELLOW', issues
    elif certainty >= 2 and hedge_count == 0:
        return 'GREEN', ['Strong, confident language']
    else:
        return 'NEUTRAL', []


# ============================================================
# PDF EXPORT
# ============================================================

def export_pdf(text, analysis, output_path, ticker=None, quarter=None,
               prior_comparison=None):
    """
    Export a PDF report with color-highlighted script
    RED = problem sentences, YELLOW = watch, GREEN = good
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    
    # Create document
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        alignment=TA_CENTER,
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
        textColor=colors.HexColor('#333333'),
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
    )
    
    # Color styles for sentences
    red_style = ParagraphStyle(
        'Red',
        parent=body_style,
        backColor=colors.HexColor('#ffcccc'),
        borderPadding=3,
    )
    
    yellow_style = ParagraphStyle(
        'Yellow',
        parent=body_style,
        backColor=colors.HexColor('#fff3cd'),
        borderPadding=3,
    )
    
    green_style = ParagraphStyle(
        'Green',
        parent=body_style,
        backColor=colors.HexColor('#d4edda'),
        borderPadding=3,
    )
    
    neutral_style = body_style
    
    # Build content
    content = []
    
    # Title
    title = "Earnings Script Analysis"
    if ticker:
        title += f" — {ticker}"
    if quarter:
        title += f" {quarter}"
    content.append(Paragraph(title, title_style))
    content.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", body_style))
    content.append(Spacer(1, 20))
    
    # Score summary
    scores = analysis['scores']
    grade_colors = {
        'A': '#28a745', 'B+': '#5cb85c', 'B': '#8bc34a',
        'C+': '#ffc107', 'C': '#ff9800', 'D': '#ff5722', 'F': '#dc3545'
    }
    grade = get_grade(scores['overall'])
    
    content.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
    
    score_data = [
        ['Overall Score', f"{int(scores['overall'])}/100 ({grade})"],
        ['Sentiment', f"{int(scores['sentiment'])}/100"],
        ['Confidence', f"{int(scores['confidence'])}/100"],
        ['Ownership', f"{int(scores['ownership'])}/100"],
        ['Clarity', f"{int(scores['clarity'])}/100"],
        ['Red Flags', f"{int(scores['red_flags'])}/100"],
    ]
    
    score_table = Table(score_data, colWidths=[2*inch, 1.5*inch])
    score_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#00d4aa')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
    ]))
    content.append(score_table)
    content.append(Spacer(1, 15))

    # Prior call comparison
    if prior_comparison and prior_comparison.get("vs_prior"):
        vs = prior_comparison["vs_prior"]
        delta = vs.get("overall_delta", 0)
        arrow = "\u2191" if delta > 0 else "\u2193" if delta < 0 else "\u2194"
        sign = "+" if delta > 0 else ""
        prior_score = prior_comparison["prior_scores"][0]["scores"]["overall"] if prior_comparison.get("prior_scores") else "?"

        content.append(Paragraph("VS. PRIOR CALL", heading_style))

        prior_data = [
            ["", "Current", vs.get("quarter", "Prior")],
            ["Overall", str(int(scores['overall'])), str(int(prior_score))],
        ]
        dim_labels = {"sentiment": "Sentiment", "confidence": "Confidence",
                      "ownership": "Ownership", "clarity": "Clarity", "red_flags": "Red Flags"}
        for dim_key, dim_label in dim_labels.items():
            cur_val = int(scores.get(dim_key, 0))
            prev_val = int(prior_comparison["prior_scores"][0]["scores"].get(dim_key, 0)) if prior_comparison.get("prior_scores") else 0
            d = cur_val - prev_val
            d_str = f"{'+' if d > 0 else ''}{d}"
            prior_data.append([dim_label, f"{cur_val} ({d_str})", str(prev_val)])

        prior_table = Table(prior_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch])
        prior_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90d9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f4ff')),
            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ]))
        content.append(prior_table)

        # Trend line if multiple quarters
        if prior_comparison.get("prior_scores") and len(prior_comparison["prior_scores"]) > 1:
            trend_parts = []
            for ps in reversed(prior_comparison["prior_scores"]):
                trend_parts.append(f"{ps['quarter']}: {ps['scores']['overall']}")
            trend_parts.append(f"Current: {int(scores['overall'])}")
            content.append(Paragraph(
                f"<i>Trend: {' → '.join(trend_parts)}</i>",
                body_style,
            ))

        content.append(Spacer(1, 15))


    # Legend
    content.append(Paragraph("COLOR LEGEND", heading_style))
    legend_data = [
        ['🔴 RED', 'Significant issues — consider revising'],
        ['🟡 YELLOW', 'Minor concerns — review if possible'],
        ['🟢 GREEN', 'Strong, confident language'],
    ]
    legend_table = Table(legend_data, colWidths=[1.2*inch, 4*inch])
    legend_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#ffcccc')),
        ('BACKGROUND', (0, 1), (0, 1), colors.HexColor('#fff3cd')),
        ('BACKGROUND', (0, 2), (0, 2), colors.HexColor('#d4edda')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    content.append(legend_table)
    content.append(Spacer(1, 25))
    
    # Highlighted script
    content.append(Paragraph("SCRIPT ANALYSIS", heading_style))
    
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    for sentence in sentences:
        if len(sentence.strip()) < 10:
            continue
            
        color, issues = classify_sentence(sentence)
        
        # Escape HTML characters
        safe_sentence = sentence.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        if issues:
            safe_sentence += f" <i>[{', '.join(issues)}]</i>"
        
        if color == 'RED':
            content.append(Paragraph(safe_sentence, red_style))
        elif color == 'YELLOW':
            content.append(Paragraph(safe_sentence, yellow_style))
        elif color == 'GREEN':
            content.append(Paragraph(safe_sentence, green_style))
        else:
            content.append(Paragraph(safe_sentence, neutral_style))
    
    # Build PDF
    doc.build(content)
    return output_path


def get_grade(score):
    if score >= 90: return 'A'
    if score >= 80: return 'B+'
    if score >= 70: return 'B'
    if score >= 60: return 'C+'
    if score >= 50: return 'C'
    if score >= 40: return 'D'
    return 'F'


# ============================================================
# WORD EXPORT WITH TRACK CHANGES
# ============================================================

def export_word(text, analysis, output_path, ticker=None, quarter=None):
    """
    Export a Word document with track changes showing suggested edits
    Uses revision marks to show original and suggested text
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Earnings Script — Suggested Revisions', 0)
    
    if ticker:
        doc.add_paragraph(f"Company: {ticker} {quarter or ''}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()
    
    # Instructions
    instructions = doc.add_paragraph()
    instructions.add_run("Instructions: ").bold = True
    instructions.add_run(
        "This document shows suggested revisions to improve your earnings script. "
        "Strikethrough text should be removed. Underlined green text shows suggested replacements. "
        "Accept or reject changes as appropriate for your communication style."
    )
    doc.add_paragraph()
    
    # Legend
    legend = doc.add_paragraph()
    legend.add_run("Legend: ").bold = True
    
    strike_run = legend.add_run("Strikethrough")
    strike_run.font.strike = True
    legend.add_run(" = Remove  |  ")
    
    add_run = legend.add_run("Underlined Green")
    add_run.font.underline = True
    add_run.font.color.rgb = RGBColor(0, 128, 0)
    legend.add_run(" = Add  |  ")
    
    highlight_run = legend.add_run("Yellow highlight")
    highlight_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    legend.add_run(" = Review")
    
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    doc.add_paragraph()
    
    # Process sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    for sentence in sentences:
        if len(sentence.strip()) < 10:
            continue
        
        color, issues = classify_sentence(sentence)
        suggested, changes = get_suggested_rewrite(sentence)
        
        para = doc.add_paragraph()
        
        if suggested and suggested != sentence:
            # Show original with strikethrough
            original_run = para.add_run(sentence)
            original_run.font.strike = True
            original_run.font.color.rgb = RGBColor(180, 0, 0)
            
            para.add_run(" → ")
            
            # Show suggested replacement
            suggested_run = para.add_run(suggested)
            suggested_run.font.underline = True
            suggested_run.font.color.rgb = RGBColor(0, 128, 0)
            
            # Add comment about what changed
            comment = para.add_run(f"  [{', '.join(changes)}]")
            comment.font.size = Pt(8)
            comment.font.italic = True
            comment.font.color.rgb = RGBColor(128, 128, 128)
            
        elif color == 'YELLOW' and issues:
            # Highlight for review
            run = para.add_run(sentence)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            comment = para.add_run(f"  [Review: {', '.join(issues)}]")
            comment.font.size = Pt(8)
            comment.font.italic = True
            comment.font.color.rgb = RGBColor(128, 128, 128)
            
        elif color == 'RED' and issues:
            # Highlight red issues
            run = para.add_run(sentence)
            run.font.highlight_color = WD_COLOR_INDEX.PINK
            
            comment = para.add_run(f"  [Issue: {', '.join(issues)}]")
            comment.font.size = Pt(8)
            comment.font.italic = True
            comment.font.color.rgb = RGBColor(180, 0, 0)
            
        else:
            # Normal text
            para.add_run(sentence)
        
        # Add space between sentences
        para.add_run(" ")
    
    # Save
    doc.save(str(output_path))
    return output_path


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    import sys
    import json
    
    print("Exporters module. Import and use export_pdf() or export_word().")
    print("\nExample:")
    print("  from exporters import export_pdf, export_word")
    print("  export_pdf(text, analysis, 'report.pdf')")
    print("  export_word(text, analysis, 'script_revisions.docx')")
