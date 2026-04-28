"""
StreetSignals.ai — FastAPI Backend
Wraps existing analysis engine in a REST API.
Integrates Claude for institutional-grade analyst Q&A.
"""

from typing import Optional, List, Dict
import asyncio
import re
import difflib
import logging
import time
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import tempfile
import os
import uuid
import json
from pathlib import Path

import httpx
import anthropic

# Existing modules — not modified
from earnings_analyzer import analyze_transcript, load_lm_dictionary
from legal_context import analyze_with_legal_context
from exporters import (
    export_pdf,
    classify_sentence,
    get_suggested_rewrite,
)
from advanced_analysis import run_advanced_analysis

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger("streetsignals")
logger.setLevel(logging.INFO)

app = FastAPI(title="StreetSignals.ai API")

# ---------------------------------------------------------------------------
# API keys from environment (set in Railway dashboard)
# ---------------------------------------------------------------------------
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
FMP_API_KEY = os.environ.get("FMP_API_KEY", "")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load Loughran-McDonald dictionary once at startup
LM_DICT = load_lm_dictionary(Path(__file__).parent / "LM_MasterDictionary.csv")

# In-memory session store — keyed by session_id
SESSIONS: dict = {}


# ---------------------------------------------------------------------------
# Overrides — fixes for scoring bugs and grading scale
# (We don't touch the sacred files; we correct in the mapping layer)
# ---------------------------------------------------------------------------

def _corrected_grade(score: float) -> str:
    """Standard academic grading scale (the sacred code's is too lenient)."""
    if score >= 93:
        return "A"
    if score >= 90:
        return "A-"
    if score >= 87:
        return "B+"
    if score >= 83:
        return "B"
    if score >= 80:
        return "B-"
    if score >= 77:
        return "C+"
    if score >= 73:
        return "C"
    if score >= 70:
        return "C-"
    if score >= 67:
        return "D+"
    if score >= 60:
        return "D"
    return "F"


def _corrected_red_flags_score(base_analysis: dict) -> int:
    """
    Fix the double-multiplication bug in the sacred code.

    The sacred code computes:
        red_flag_pct = (count / total) * 100   # already a percentage, e.g. 2.5
    Then scores it as:
        scores['red_flags'] = 100 - (red_flag_pct * 100)   # multiplies again!

    So even 1% red flags → 100 - 100 = 0.  We fix that here.
    A score of 100 means NO red flags (good), 0 means lots of red flags (bad).
    """
    deception = base_analysis.get("deception", {})
    red_flag_pct = deception.get("red_flag_pct", 0)
    # red_flag_pct is already multiplied by 100, so treat it as a percentage directly
    # e.g. red_flag_pct = 2.5 means 2.5% of words are red flags
    # Scale: 0% → 100, 5% → 50, 10%+ → 0
    score = max(0, min(100, round(100 - (red_flag_pct * 10))))
    return score


def _corrected_scores(base_analysis: dict) -> dict:
    """Recompute scores with the red_flags bug fixed and new grade scale."""
    raw_scores = base_analysis.get("scores", {})
    fixed_red = _corrected_red_flags_score(base_analysis)

    sentiment = round(raw_scores.get("sentiment", 0))
    confidence = round(raw_scores.get("confidence", 0))
    ownership = round(raw_scores.get("ownership", 0))
    clarity = round(raw_scores.get("clarity", 0))

    weights = {
        "sentiment": 0.25,
        "confidence": 0.25,
        "ownership": 0.15,
        "clarity": 0.15,
        "red_flags": 0.20,
    }
    overall = round(
        sentiment * weights["sentiment"]
        + confidence * weights["confidence"]
        + ownership * weights["ownership"]
        + clarity * weights["clarity"]
        + fixed_red * weights["red_flags"]
    )

    return {
        "overall": overall,
        "grade": _corrected_grade(overall),
        "sentiment": sentiment,
        "confidence": confidence,
        "ownership": ownership,
        "clarity": clarity,
        "red_flags": fixed_red,
    }


# ---------------------------------------------------------------------------
# Rewrite helpers — ensure every flagged issue gets a suggested rewrite
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Claude-powered rewrites — context-aware, respects legal language
# ---------------------------------------------------------------------------

_REWRITE_SYSTEM_PROMPT = """You are an expert IR (Investor Relations) editor reviewing an earnings call script.

Your job: For each flagged sentence, propose a concrete rewrite that addresses the flagged issues while respecting legal and contextual constraints.

CRITICAL RULES:
1. NEVER modify safe harbor statements, forward-looking statement disclaimers, or legal boilerplate. For these, propose a rewrite that adds context or strengthens the surrounding language while keeping the legally required words intact.
2. NEVER change cautionary risk disclosures into definitive claims. Instead, tighten the language to sound more deliberate and less uncertain.
3. Preserve grammatical correctness — every rewrite must read naturally.
4. Preserve the original meaning. Don't turn uncertain statements into definitive claims.
5. Make minimal changes — change as few words as possible.
6. EVERY sentence MUST get a rewrite that improves it, even if the improvement is subtle. No sentence should be returned unchanged.

STRATEGIES FOR DIFFERENT ISSUE TYPES:
- Hedging (unnecessary): Remove filler words like "somewhat", "relatively", "generally" when they add no meaning.
- Hedging (legal/cautionary): Keep the hedging words but tighten the sentence structure, remove redundancy, or add specificity.
- Weak verbs: "We hope" → "We expect", "We feel" → "We are confident", "We think" → "We believe"
- Vague language: Add specificity or restructure for directness.
- Safe harbor boilerplate: Rewrite surrounding context to be more direct while preserving required legal language verbatim.

EXAMPLES:
- "We are cautiously optimistic about growth" → "We are optimistic about growth"
- "We hope to see improvement" → "We expect to see improvement"
- "We believe our AI systems may be contributing to engagement" → "We believe our AI systems are contributing to engagement"
- "Actual results may differ materially from those anticipated." → "Actual results may differ materially from those anticipated, and we encourage investors to review the risk factors in our SEC filings."

Return your rewrites in this EXACT numbered format — one rewrite per line:
[0] The improved first sentence
[1] The improved second sentence

CRITICAL: Include ALL indices from the input. Every sentence MUST appear. Do NOT add any explanations, commentary, or extra text — ONLY the numbered rewrites."""


async def _generate_rewrites_with_claude(flagged_sentences: list) -> Optional[dict]:
    """
    Send all flagged sentences to Claude in one batch for context-aware rewriting.
    Returns a dict mapping sentence text → rewrite (or same text if no change).
    """
    if not ANTHROPIC_API_KEY or not flagged_sentences:
        logger.info(f"Skipping Claude rewrites: key={'set' if ANTHROPIC_API_KEY else 'MISSING'}, "
                     f"sentences={len(flagged_sentences) if flagged_sentences else 0}")
        return None

    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        logger.info(f"Sending {len(flagged_sentences)} flagged sentences to Claude for rewriting")

        # Build the prompt with numbered sentences
        parts = [
            "Here are the flagged sentences from an earnings call script.",
            "For EVERY sentence, propose a concrete rewrite that addresses the flagged issues.",
            "Do NOT return any sentence unchanged.\n",
        ]
        for i, item in enumerate(flagged_sentences):
            parts.append(f"[{i}] Issues: {', '.join(item['issues'])}")
            parts.append(f"    Sentence: {item['sentence']}")
            parts.append("")

        response = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=16000,
            system=_REWRITE_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": "\n".join(parts)}],
        )

        response_text = response.content[0].text.strip()
        logger.info(f"Claude rewrite response: {len(response_text)} chars, "
                     f"stop_reason={response.stop_reason}")

        # Parse numbered format: [idx] rewrite text
        # Far more robust than JSON — handles preamble, code fences, partial output
        result = {}
        pattern = re.compile(r'^\[(\d+)\]\s*(.+)$', re.MULTILINE)
        for match in pattern.finditer(response_text):
            idx = int(match.group(1))
            rewrite_text = match.group(2).strip()
            if 0 <= idx < len(flagged_sentences):
                original = flagged_sentences[idx]["sentence"]
                result[original] = rewrite_text
                norm_key = ' '.join(original.split())
                if norm_key != original:
                    result[norm_key] = rewrite_text

        logger.info(f"Claude: {len(result)} rewrites parsed for {len(flagged_sentences)} sentences "
                     f"(stop_reason={response.stop_reason})")

        if not result:
            logger.error(f"0 rewrites parsed! Response starts with: {response_text[:500]}")

        return result if result else None

    except Exception as e:
        logger.error(f"Claude rewrite generation failed: {type(e).__name__}: {e}")
        return None


def _generate_rewrite_fallback(sentence: str, issues: list) -> str:
    """Simple fallback rewrite when Claude is unavailable. Returns original if no good rewrite."""
    rewrite, changes = get_suggested_rewrite(sentence)
    if rewrite and rewrite != sentence:
        return rewrite
    return sentence


# ---------------------------------------------------------------------------
# Claude-powered negative interpretations + guidance extraction (3rd call)
# ---------------------------------------------------------------------------

_ANALYSIS_SYSTEM_PROMPT = """You are a senior sell-side equity research analyst who specializes in identifying language in earnings call scripts that could be spun negatively by bearish analysts or short sellers.

You will analyze an earnings call script and provide FOUR analyses:

PART 1 — NEGATIVE INTERPRETATIONS:
Identify 5-15 passages in the script that a bearish analyst could use to build a negative narrative. For each:
- Quote the problematic text (exact words from the script)
- Explain how a bear analyst would spin it
- Suggest a rewrite that neutralizes the negative spin
- Rate severity: "high" (actively harmful), "medium" (concerning), "low" (minor risk)
- Categorize: hedging_language, vague_commitments, mixed_messaging, defensiveness, omission_signal, metric_avoidance, blame_shifting, over_promising, vague_guidance, missing_guidance, guidance_gap

PART 2 — GUIDANCE EXTRACTION:
Extract any forward-looking guidance, outlook, or targets from the script. Only include ACTUAL guidance — where management explicitly provides a forecast, target, outlook, or range for a future metric. Do NOT include:
- Past results or historical numbers
- General commentary like "we're optimistic" without specific targets
- Analyst questions (only management's stated guidance)
For each metric with guidance:
- Metric name (e.g., "Revenue", "Operating Margin", "Operating Cash Flow", "EPS", "Gross Margin", "CapEx")
- The specific guidance value or range given (e.g., "$4.2B - $4.4B", "25-27%", "mid-single-digit growth")
- Whether it's quantified (true = specific number/range/percentage, false = directional only like "growth" or "improvement")
- The exact quote from the script where this guidance was stated

If NO forward-looking guidance is provided in the script, return an empty array for guidance_metrics.

IMPORTANT — GUIDANCE-INFORMED NEGATIVE INTERPRETATIONS:
After extracting guidance metrics above, use that analysis to generate ADDITIONAL negative interpretations in PART 1 with these categories:
- "vague_guidance": Flag any guidance that uses hedging language ("we anticipate", "approximately", "in the neighborhood of") instead of specific ranges or numbers. Quote the vague language and explain how a sophisticated analyst would press for specifics.
- "missing_guidance": Identify key financial metrics where analysts would expect forward guidance but the script provides NONE (e.g., revenue, margins, EPS, CapEx, free cash flow). Explain what would be conspicuously absent to a buy-side analyst.
- "guidance_gap": Flag guidance with unusually wide ranges, lack of year-over-year comparison, or purely directional language ("improvement", "growth") without quantifying. Explain how this creates uncertainty that the market will discount.
These guidance-related findings MUST appear as entries in the "negative_interpretations" array with the appropriate category. Do not put them in guidance_metrics.

PART 3 — LITIGATION RISK VALIDATION:
You will be given a list of pattern-matched litigation risk findings. For each, determine if it's a REAL litigation risk or a false positive. Only keep findings that represent genuine legal exposure risks (securities fraud language, misleading claims, material omissions, regulatory issues). Remove false positives where:
- The flagged text is actually reasonable/normal business language
- The "point estimate without range" finding doesn't represent actual forward guidance
- The finding is about past results, not forward-looking statements
Return only validated findings with specific, actionable recommendations tied to the actual text.
CRITICAL: For each validated finding, you MUST provide a "suggested_rewrite" — a revised version of the original_text that mitigates the legal risk while preserving the intended business message. This rewrite will appear as tracked changes in the Word document.

PART 4 — ACTIVIST TRIGGER VALIDATION:
You will be given a list of pattern-matched activist vulnerability triggers. For each, determine if it's a REAL activist concern or noise. Only keep triggers that an activist investor would actually seize upon (capital allocation issues, governance concerns, margin underperformance, strategic drift). For each validated trigger:
- Explain the specific activist narrative
- Provide a defense suggestion specific to the actual concern
CRITICAL: For each validated trigger, you MUST provide a "suggested_rewrite" — a revised version of the original_text that defuses the activist concern while preserving the intended business message. This rewrite will appear as tracked changes in the Word document.

Emit your analysis by calling the `emit_analysis` tool with all four fields (negative_interpretations, guidance_metrics, litigation_findings, activist_triggers). Every string field must contain ONLY the value itself — do not append parenthetical clarifications like '(in context: ...)' to string values; if context is needed, put it in the appropriate separate field.

REWRITE UNIQUENESS: Each suggested_rewrite text must be distinct from every other suggested_rewrite in the response. When two adjacent original sentences make essentially the same point (e.g. "This is not a churn story." followed by "This is a timing-of-expansion story."), do NOT skip flagging them — instead, write a DIFFERENT, sentence-specific rewrite for each one. Each rewrite should preserve and address the unique angle of its original sentence. Both flags are valuable; just make sure the two suggested_rewrites are different sentences."""


# Tool schema — forces Claude to return structured JSON that can't be malformed.
# Every free-form field is a plain string; context goes in its own field so the
# model can't smuggle extra structure inside a string and break parsing.
_ANALYSIS_TOOL = {
    "name": "emit_analysis",
    "description": "Emit the earnings call analysis in structured form. Call this exactly once with all four arrays populated.",
    "input_schema": {
        "type": "object",
        "properties": {
            "negative_interpretations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {"type": "string"},
                        "original_text": {"type": "string", "description": "Exact quote from the script — no parentheticals, no commentary."},
                        "negative_spin": {"type": "string"},
                        "suggested_rewrite": {"type": "string"},
                        "severity": {"type": "string", "enum": ["low", "medium", "high"]},
                    },
                    "required": ["category", "original_text", "negative_spin", "suggested_rewrite", "severity"],
                },
            },
            "guidance_metrics": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "metric": {"type": "string"},
                        "value": {"type": "string"},
                        "quantified": {"type": "boolean"},
                        "quote": {"type": "string"},
                    },
                    "required": ["metric", "value", "quantified", "quote"],
                },
            },
            "litigation_findings": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "issue": {"type": "string"},
                        "detail": {"type": "string"},
                        "original_text": {"type": "string"},
                        "suggested_rewrite": {"type": "string"},
                        "recommendation": {"type": "string"},
                        "severity": {"type": "string", "enum": ["Low", "Medium", "High"]},
                    },
                    "required": ["issue", "detail", "original_text", "suggested_rewrite", "recommendation", "severity"],
                },
            },
            "activist_triggers": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {"type": "string"},
                        "original_text": {"type": "string"},
                        "suggested_rewrite": {"type": "string"},
                        "activist_narrative": {"type": "string"},
                        "defense_suggestion": {"type": "string"},
                        "severity": {"type": "string", "enum": ["low", "medium", "high"]},
                    },
                    "required": ["category", "original_text", "suggested_rewrite", "activist_narrative", "defense_suggestion", "severity"],
                },
            },
        },
        "required": ["negative_interpretations", "guidance_metrics", "litigation_findings", "activist_triggers"],
    },
}


async def _generate_analysis_with_claude(
    script: str,
    base_negative_interps: list,
    base_litigation_findings: list = None,
    base_activist_triggers: list = None,
) -> Optional[dict]:
    """
    3rd parallel Claude call: generates context-aware negative interpretations,
    extracts actual guidance metrics, validates litigation findings, and
    validates activist triggers. Claude is the source of truth for all 4 advanced tabs.
    """
    if not ANTHROPIC_API_KEY:
        return None

    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        parts = []

        # Include base pattern-matched negative interps as context
        if base_negative_interps:
            parts.append("PATTERN-BASED NEGATIVE INTERPRETATION FINDINGS (for context — build on these, don't just repeat them):")
            for item in base_negative_interps[:10]:
                parts.append(f"- [{item.get('severity', 'medium')}] \"{item.get('matched_text', '')}\" — {item.get('interpretation', '')}")
            parts.append("")

        # Include base litigation findings for validation
        if base_litigation_findings:
            parts.append("PATTERN-BASED LITIGATION RISK FINDINGS (validate these — remove false positives, keep only real risks):")
            for f in base_litigation_findings[:15]:
                parts.append(f"- [{f.get('severity', 'Medium')}] {f.get('issue', '')} — \"{f.get('sentence', '')[:150]}\"")
                if f.get('detail'):
                    parts.append(f"  Detail: {f.get('detail', '')[:200]}")
            parts.append("")

        # Include base activist triggers for validation
        if base_activist_triggers:
            parts.append("PATTERN-BASED ACTIVIST TRIGGER FINDINGS (validate these — remove noise, keep only real activist concerns):")
            for t in base_activist_triggers[:15]:
                parts.append(f"- [{t.get('severity', 'medium')}] {t.get('trigger', '')} — \"{t.get('sentence', '')[:150]}\"")
                if t.get('concern'):
                    parts.append(f"  Concern: {t.get('concern', '')[:200]}")
            parts.append("")

        # The script
        parts.append("EARNINGS CALL SCRIPT:")
        parts.append("=" * 60)
        script_words = script.split()
        if len(script_words) > 8000:
            parts.append(" ".join(script_words[:8000]))
            parts.append(f"[...truncated from {len(script_words)} words]")
        else:
            parts.append(script)
        parts.append("=" * 60)

        parts.append("\nAnalyze this script and return ALL FOUR parts as specified: negative interpretations, guidance extraction, litigation findings validation, and activist trigger validation.")

        response = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            system=_ANALYSIS_SYSTEM_PROMPT,
            tools=[_ANALYSIS_TOOL],
            tool_choice={"type": "tool", "name": "emit_analysis"},
            messages=[{"role": "user", "content": "\n".join(parts)}],
        )

        tool_use = next((block for block in response.content if block.type == "tool_use"), None)
        if tool_use is None:
            logger.error(f"Claude analysis: no tool_use block in response (stop_reason={response.stop_reason})")
            return None

        data = tool_use.input

        neg_interps = data.get("negative_interpretations", [])
        guidance = data.get("guidance_metrics", [])
        lit_findings = data.get("litigation_findings", [])
        activist = data.get("activist_triggers", [])

        # Deduplicate by suggested_rewrite — Claude sometimes returns the same
        # rewrite for two adjacent sentences making the same point (e.g. "This
        # is not a churn story." and "This is a timing-of-expansion story."
        # both becoming the same neutral statement). Showing the same
        # replacement twice in the doc looks broken; keep only the first.
        def _dedupe_by_rewrite(items, label):
            seen = set()
            out = []
            for item in items:
                if not isinstance(item, dict):
                    continue
                rw = (item.get("suggested_rewrite") or "").strip()
                if rw and rw in seen:
                    logger.info(
                        f"Dropping duplicate {label} rewrite for "
                        f"original={item.get('original_text', '')[:60]!r}"
                    )
                    continue
                if rw:
                    seen.add(rw)
                out.append(item)
            return out

        neg_interps = _dedupe_by_rewrite(neg_interps, "neg_interp")
        lit_findings = _dedupe_by_rewrite(lit_findings, "litigation")
        activist = _dedupe_by_rewrite(activist, "activist")

        logger.info(
            f"Claude analysis: {len(neg_interps)} neg interps, "
            f"{len(guidance)} guidance metrics, "
            f"{len(lit_findings)} litigation findings, "
            f"{len(activist)} activist triggers"
        )

        return {
            "negative_interpretations": neg_interps,
            "guidance_metrics": guidance,
            "litigation_findings": lit_findings,
            "activist_triggers": activist,
        }

    except json.JSONDecodeError as e:
        # Structured tool-use should eliminate this class of error. If it recurs
        # the error will point at Anthropic's own serialization, which is a bug
        # on their side — log what we can and fall back.
        logger.error(
            f"Claude analysis JSON parse error: {e} — unexpected with tool-use mode"
        )
        return None
    except Exception as e:
        logger.error(f"Claude analysis generation failed: {type(e).__name__}: {e}")
        return None


# ---------------------------------------------------------------------------
# 4th parallel Claude call: Bull/Bear Case Defense
# ---------------------------------------------------------------------------

_BULL_BEAR_SYSTEM_PROMPT = """You are a senior equity research analyst at JP Morgan with deep sector coverage.

Your task has two parts:

PART 1 — BULL/BEAR CASES:
Identify the 3-5 strongest bull cases and 3-5 strongest bear cases for the given company.
These should reflect the current Wall Street consensus debate — the real arguments institutional investors are making. For each case, provide a one-sentence thesis and a 2-3 sentence explanation.

PART 2 — SCRIPT REWRITES:
Review the earnings call script. For each bear case, find sentences that:
(a) ignore the bear case when they should address it head-on,
(b) inadvertently give credence to the bear narrative, or
(c) miss an opportunity to defuse it with data or framing.

For each bull case, find sentences that could more strongly reinforce the bull thesis with specific language improvements.

For each rewrite, quote the EXACT original sentence from the script and provide an improved version. The rewrite should sound natural for an executive delivering an earnings call — authoritative, specific, and confident without being legally reckless.

Emit your analysis via the `emit_bull_bear` tool."""


_BULL_BEAR_TOOL = {
    "name": "emit_bull_bear",
    "description": "Emit bull/bear cases and proposed script rewrites. Call exactly once.",
    "input_schema": {
        "type": "object",
        "properties": {
            "bull_cases": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "thesis": {"type": "string"},
                        "explanation": {"type": "string"},
                    },
                    "required": ["thesis", "explanation"],
                },
            },
            "bear_cases": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "thesis": {"type": "string"},
                        "explanation": {"type": "string"},
                    },
                    "required": ["thesis", "explanation"],
                },
            },
            "rewrites": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "original_text": {"type": "string"},
                        "suggested_rewrite": {"type": "string"},
                        "case_type": {"type": "string", "enum": ["bull", "bear"]},
                        "case_thesis": {"type": "string"},
                        "rationale": {"type": "string"},
                    },
                    "required": ["original_text", "suggested_rewrite", "case_type", "case_thesis", "rationale"],
                },
            },
        },
        "required": ["bull_cases", "bear_cases", "rewrites"],
    },
}


async def _generate_bull_bear_with_claude(
    script: str,
    ticker: str,
) -> Optional[dict]:
    """
    4th parallel Claude call: identifies bull/bear investment cases for the
    company and proposes sentence-level rewrites to the earnings script that
    address bear concerns and reinforce bull arguments.
    Only runs when a ticker is provided.
    """
    if not ANTHROPIC_API_KEY or not ticker:
        return None

    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        script_words = script.split()
        if len(script_words) > 8000:
            script_text = " ".join(script_words[:8000])
            script_text += f"\n[...truncated from {len(script_words)} words]"
        else:
            script_text = script

        user_message = (
            f"COMPANY: {ticker.upper()}\n\n"
            f"EARNINGS CALL SCRIPT:\n{'=' * 60}\n"
            f"{script_text}\n{'=' * 60}\n\n"
            f"Analyze the bull and bear cases for {ticker.upper()} and propose "
            f"rewrites to this earnings script as specified."
        )

        response = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8000,
            system=_BULL_BEAR_SYSTEM_PROMPT,
            tools=[_BULL_BEAR_TOOL],
            tool_choice={"type": "tool", "name": "emit_bull_bear"},
            messages=[{"role": "user", "content": user_message}],
        )

        tool_use = next((block for block in response.content if block.type == "tool_use"), None)
        if tool_use is None:
            logger.error(f"Claude bull/bear: no tool_use block (stop_reason={response.stop_reason})")
            return None

        data = tool_use.input

        # Defensive: despite the tool schema declaring objects, the model can
        # still emit string items under load/edge cases. Drop anything that
        # isn't a dict so downstream code doesn't crash on rw.get().
        def _only_dicts(items, label):
            if not isinstance(items, list):
                logger.warning(f"Claude bull/bear: {label} was {type(items).__name__}, coercing to []")
                return []
            clean = [x for x in items if isinstance(x, dict)]
            if len(clean) != len(items):
                logger.warning(
                    f"Claude bull/bear: dropped {len(items) - len(clean)} non-dict entries from {label}"
                )
            return clean

        bull_cases = _only_dicts(data.get("bull_cases", []), "bull_cases")
        bear_cases = _only_dicts(data.get("bear_cases", []), "bear_cases")
        rewrites = _only_dicts(data.get("rewrites", []), "rewrites")

        logger.info(
            f"Claude bull/bear for {ticker}: {len(bull_cases)} bull cases, "
            f"{len(bear_cases)} bear cases, {len(rewrites)} rewrites"
        )

        return {
            "bull_cases": bull_cases,
            "bear_cases": bear_cases,
            "rewrites": rewrites,
        }

    except json.JSONDecodeError as e:
        logger.error(f"Claude bull/bear JSON parse error: {e}")
        return None
    except Exception as e:
        logger.error(f"Claude bull/bear generation failed: {type(e).__name__}: {e}")
        return None


# ---------------------------------------------------------------------------
# FMP consensus estimates fetch
# ---------------------------------------------------------------------------

async def _fetch_consensus_estimates(ticker: str) -> Optional[dict]:
    """Fetch analyst consensus estimates from FMP. Returns None if unavailable."""
    if not FMP_API_KEY or not ticker:
        return None

    ticker = ticker.strip().upper()

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(
                "https://financialmodelingprep.com/stable/analyst-estimates",
                params={"symbol": ticker, "apikey": FMP_API_KEY, "limit": 1},
            )
            if resp.status_code != 200:
                logger.warning(f"FMP consensus fetch returned {resp.status_code} for {ticker}")
                return None

            data = resp.json()
            if not data or not isinstance(data, list) or len(data) == 0:
                return None

            est = data[0]
            return {
                "revenue": est.get("estimatedRevenueAvg"),
                "earnings": est.get("estimatedEpsAvg"),
                "net_income": est.get("estimatedNetIncomeAvg"),
                "ebitda": est.get("estimatedEbitdaAvg"),
                "sga": est.get("estimatedSgaExpenseAvg"),
                "date": est.get("date"),
                "source": "fmp",
            }
    except Exception as e:
        logger.warning(f"FMP consensus fetch failed for {ticker}: {type(e).__name__}: {e}")
        return None


# ---------------------------------------------------------------------------
# Mapping helpers — transform existing output into the spec's JSON shape
# ---------------------------------------------------------------------------

def _build_flagged_issues(
    text: str,
    base_analysis: dict,
    claude_rewrites: Optional[dict] = None,
    claude_bull_bear: Optional[dict] = None,
) -> list:
    """
    Enrich the base flagged_passages with colour and suggested rewrites.
    Uses Claude rewrites when available, falls back to sacred rewriter.
    Also merges bull/bear defense rewrites as additional flagged issues.
    """
    issues = []
    seen_sentences = set()  # Track sentences already flagged (for dedup with bull/bear)

    for fp in base_analysis.get("flagged_passages", []):
        sentence = fp["sentence"]
        color, _ = classify_sentence(sentence)
        seen_sentences.add(' '.join(sentence.split()))

        # Use Claude rewrite if available (try exact match, then normalized)
        rewrite = None
        if claude_rewrites:
            rewrite = claude_rewrites.get(sentence) or claude_rewrites.get(' '.join(sentence.split()))
        if not rewrite or rewrite == sentence:
            rewrite = _generate_rewrite_fallback(sentence, fp["issues"])

        # Only include rewrite if it's actually different
        issues.append(
            {
                "sentence": sentence,
                "color": color,
                "issues": fp["issues"],
                "suggested_rewrite": rewrite if rewrite != sentence else None,
            }
        )

    # Append bull/bear defense rewrites as additional flagged issues
    if claude_bull_bear and "rewrites" in claude_bull_bear:
        for rw in claude_bull_bear["rewrites"]:
            if not isinstance(rw, dict):
                continue
            orig = rw.get("original_text", "")
            rewrite = rw.get("suggested_rewrite", "")
            case_type = rw.get("case_type", "bear").title()
            thesis = rw.get("case_thesis", "")
            rationale = rw.get("rationale", "")

            if not orig or not rewrite or rewrite == orig:
                continue

            norm_key = ' '.join(orig.split())

            # If this sentence is already flagged, add the bull/bear tag to it
            # rather than creating a duplicate entry
            already_flagged = False
            if norm_key in seen_sentences:
                for issue in issues:
                    if ' '.join(issue["sentence"].split()) == norm_key:
                        tag = f"Bull/Bear Defense: {case_type} Case"
                        issue["issues"] = issue["issues"] + [tag]
                        # Bull/bear rewrite takes priority (addresses narrative + issue)
                        issue["suggested_rewrite"] = rewrite
                        if rationale:
                            issue["bull_bear_rationale"] = rationale
                        already_flagged = True
                        break

            if not already_flagged:
                thesis_short = thesis[:80] if thesis else ""
                issues.append({
                    "sentence": orig,
                    "color": "YELLOW",
                    "issues": [f"Bull/Bear Defense: {case_type} Case — {thesis_short}"],
                    "suggested_rewrite": rewrite,
                    "source": "bull_bear",
                    "case_type": rw.get("case_type", "bear"),
                    "case_thesis": thesis,
                    "rationale": rationale,
                })

    return issues


def _build_analyst_qa_fallback(advanced: dict) -> dict:
    """Template-based Q&A from the sacred code — used as fallback."""
    questions = advanced.get("analyst_questions", [])
    answers = advanced.get("proposed_answers", [])
    return {
        "total_questions": len(questions),
        "questions": questions,
        "answers": answers,
    }


# ---------------------------------------------------------------------------
# FMP transcript fetching — auto-pull prior earnings call transcripts
# ---------------------------------------------------------------------------

async def _fetch_prior_transcripts(ticker: str, num_quarters: int = 4) -> List[dict]:
    """Fetch the last N quarterly earnings call transcripts from Financial Modeling Prep."""
    if not FMP_API_KEY or not ticker:
        return []

    ticker = ticker.strip().upper()

    # Calculate the last N quarters
    now = datetime.now()
    quarters = []
    year, quarter = now.year, (now.month - 1) // 3  # 0-indexed current quarter
    # Start from the most recently completed quarter
    for _ in range(num_quarters):
        if quarter == 0:
            quarter = 4
            year -= 1
        quarters.append((year, quarter))
        quarter -= 1

    transcripts = []
    async with httpx.AsyncClient(timeout=15.0) as client:
        for yr, qtr in quarters:
            try:
                resp = await client.get(
                    "https://financialmodelingprep.com/stable/earning-call-transcript",
                    params={"symbol": ticker, "quarter": qtr, "year": yr, "apikey": FMP_API_KEY},
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if data and isinstance(data, list) and len(data) > 0:
                        content = data[0].get("content", "")
                        if content and len(content) > 200:
                            transcripts.append({
                                "quarter": f"Q{qtr} {yr}",
                                "content": content,
                            })
            except Exception as e:
                logger.warning(f"FMP fetch failed for {ticker} Q{qtr} {yr}: {e}")
                continue

    return transcripts


# ---------------------------------------------------------------------------
# Claude-powered analyst Q&A generation
# ---------------------------------------------------------------------------

_QA_SYSTEM_PROMPT = """You are a panel of senior sell-side equity research analysts preparing for an upcoming earnings call Q&A session. You have deep expertise in financial analysis and have covered this company for years.

Your task: Based on the company's prepared remarks (and optionally prior call transcripts), generate the toughest, most probing questions that analysts will ask — and draft proposed management responses.

QUESTION GUIDELINES:
- Ask specific, data-driven questions (reference actual numbers and claims from the script)
- Follow up on commitments or guidance from prior calls when available
- Probe hedging language, vague promises, or evasive areas
- Ask about topics conspicuously absent from the prepared remarks
- Challenge overly optimistic framing with requests for specifics
- Identify metrics where forward guidance is conspicuously absent or vague — probe for specific numbers, ranges, and timelines
- If guidance uses hedging language ("approximately", "in the neighborhood of") without specific ranges, ask for quantification
- Compare any stated guidance against what a sophisticated analyst would expect — ask about guidance relative to prior period performance or prior commitments
- Each question should feel like it comes from a real analyst who does this for a living

PRIOR CALL Q&A ANALYSIS (critical when prior transcripts are available):
- Pay close attention to the Q&A sections of prior earnings calls — these are the questions analysts ACTUALLY asked
- Identify questions where management gave evasive, deflective, or incomplete answers — analysts WILL follow up on these
- Track specific commitments management made in prior Q&A responses (timelines, targets, milestones) — ask about their status
- Note recurring analyst concerns across multiple quarters — persistent themes signal unresolved issues the Street cares about
- If management promised to "provide more detail next quarter" or "get back to you on that," hold them to it
- Identify metrics or topics analysts probed in prior Q&A that are NOT addressed in the current prepared remarks — these will be asked again

ANSWER GUIDELINES:
- Responses should be what a well-prepared IR/management team would say
- Be direct, confident, and data-driven — avoid corporate waffle
- Reference specific metrics, timelines, and commitments where possible
- Flag areas where management should have data ready but shouldn't speculate
- When prior Q&A is available, proactively address known analyst concerns rather than waiting to be asked

Emit your Q&A via the `emit_qa` tool. Generate 8-10 questions ordered by likelihood of being asked. The questions and answers arrays must have the same length — one answer per question, same index."""


_QA_TOOL = {
    "name": "emit_qa",
    "description": "Emit the predicted analyst Q&A. Call exactly once with matched questions and answers arrays (same length, one answer per question).",
    "input_schema": {
        "type": "object",
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "question": {"type": "string"},
                        "topic": {"type": "string"},
                        "confidence": {"type": "number", "minimum": 0, "maximum": 1},
                    },
                    "required": ["question", "topic", "confidence"],
                },
            },
            "answers": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "answer_strategy": {"type": "string"},
                        "proposed_answer": {"type": "string"},
                        "key_data_points": {"type": "array", "items": {"type": "string"}},
                        "caution_notes": {"type": "string"},
                    },
                    "required": ["answer_strategy", "proposed_answer", "key_data_points", "caution_notes"],
                },
            },
        },
        "required": ["questions", "answers"],
    },
}


async def _generate_qa_with_claude(
    current_script: str,
    prior_transcripts: List[dict],
    base_analysis: dict,
) -> Optional[dict]:
    """Generate institutional-grade analyst Q&A using Claude API."""
    if not ANTHROPIC_API_KEY:
        return None

    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        # Build the user prompt
        parts = []

        # Include prior transcripts if available
        if prior_transcripts:
            parts.append("PRIOR EARNINGS CALL TRANSCRIPTS (most recent first):")
            parts.append("=" * 60)
            for t in prior_transcripts:
                # Cap each transcript to ~5000 words to manage token usage
                words = t["content"].split()
                truncated = " ".join(words[:5000])
                parts.append(f"\n--- {t['quarter']} EARNINGS CALL ---")
                parts.append(truncated)
                if len(words) > 5000:
                    parts.append(f"[...truncated from {len(words)} words]")
            parts.append("=" * 60)
            parts.append("")

        # Include analysis context (scores and flagged issues help Claude target weak spots)
        scores = base_analysis.get("scores", {})
        parts.append("ANALYSIS CONTEXT (weak areas to probe):")
        parts.append(f"- Sentiment score: {scores.get('sentiment', 'N/A')}/100")
        parts.append(f"- Confidence score: {scores.get('confidence', 'N/A')}/100")
        parts.append(f"- Clarity score: {scores.get('clarity', 'N/A')}/100")
        flagged = base_analysis.get("flagged_passages", [])
        if flagged:
            parts.append(f"- {len(flagged)} passages flagged for hedging/vagueness")
        parts.append("")

        # The current script
        parts.append("CURRENT PREPARED REMARKS FOR UPCOMING CALL:")
        parts.append("=" * 60)
        # Cap at ~8000 words
        script_words = current_script.split()
        if len(script_words) > 8000:
            parts.append(" ".join(script_words[:8000]))
            parts.append(f"[...truncated from {len(script_words)} words]")
        else:
            parts.append(current_script)
        parts.append("=" * 60)

        parts.append("")
        if prior_transcripts:
            parts.append(
                f"You have {len(prior_transcripts)} prior call transcripts above (including their Q&A sections). "
                "Reference specific prior commitments, guidance changes, and analyst concerns from those calls. "
                "Pay special attention to the Q&A sections — identify questions analysts asked that management "
                "deflected or gave incomplete answers to, and commitments that need follow-up this quarter."
            )
        else:
            parts.append(
                "No prior transcripts available. Focus on the current prepared remarks — "
                "identify weak spots, hedging, vagueness, missing topics, and overly optimistic claims."
            )

        parts.append("\nGenerate 8-10 analyst questions with proposed management responses.")

        user_prompt = "\n".join(parts)

        # Call Claude (API data is never used for training per Anthropic's API terms)
        response = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=12000,
            system=_QA_SYSTEM_PROMPT,
            tools=[_QA_TOOL],
            tool_choice={"type": "tool", "name": "emit_qa"},
            messages=[{"role": "user", "content": user_prompt}],
        )

        if response.stop_reason == "max_tokens":
            logger.warning(
                "Claude Q&A response hit max_tokens ceiling — raise max_tokens if this recurs"
            )

        tool_use = next((block for block in response.content if block.type == "tool_use"), None)
        if tool_use is None:
            logger.error(f"Claude Q&A: no tool_use block in response (stop_reason={response.stop_reason})")
            return None

        qa_data = tool_use.input
        questions = qa_data.get("questions", [])
        answers = qa_data.get("answers", [])

        # Validate structure
        if not questions or not answers:
            logger.warning("Claude returned empty Q&A")
            return None

        return {
            "total_questions": len(questions),
            "questions": questions,
            "answers": answers,
            "source": "claude",
            "prior_calls_used": len(prior_transcripts),
        }

    except json.JSONDecodeError as e:
        logger.error(f"Claude Q&A JSON parse error: {e}")
        return None
    except anthropic.APIError as e:
        logger.error(f"Claude API error: {e}")
        return None
    except Exception as e:
        logger.error(f"Claude Q&A generation failed: {e}")
        return None


def _build_litigation(advanced: dict, claude_litigation: Optional[list] = None) -> dict:
    lit = advanced.get("litigation_risk", {})

    # Use Claude-validated findings (filtered for real risks) if available
    if claude_litigation is not None:
        findings = claude_litigation
    else:
        # Fallback to sacred code findings (may contain false positives)
        findings = lit.get("findings", [])

    # Recalculate risk level based on validated findings
    if claude_litigation is not None:
        if len(findings) == 0:
            risk_level = "Low"
        elif any(f.get("severity", "").lower() == "high" for f in findings):
            risk_level = "High"
        elif len(findings) >= 3:
            risk_level = "Medium"
        else:
            risk_level = "Low"
    else:
        risk_level = lit.get("risk_level", "Low")

    return {
        "risk_level": risk_level,
        "risk_score": lit.get("risk_score", 0),
        "has_safe_harbor": lit.get("has_safe_harbor", False),
        "findings": findings,
    }


def _build_activist(advanced: dict, claude_activist: Optional[list] = None) -> dict:
    act = advanced.get("activist_triggers", {})

    # Use Claude-validated triggers if available (filtered for real concerns)
    if claude_activist is not None:
        # Claude already returns the right format
        mapped = claude_activist
    else:
        # Fallback: map sacred code fields → frontend expected fields
        raw_triggers = act.get("triggers", [])
        mapped = []
        for t in raw_triggers:
            category = (t.get("trigger", "") or "").replace("_", " ").title()
            # Skip triggers with empty category — they're noise
            if not category.strip():
                continue
            concern = t.get("concern", "")
            mapped.append({
                "category": category,
                "trigger_type": category,
                "original_text": t.get("sentence", t.get("matched_text", "")),
                "matched_text": t.get("matched_text", ""),
                "activist_narrative": t.get("activist_angle", ""),
                "defense_suggestion": (
                    f"Consider proactively addressing this by providing specific data points "
                    f"and timeline commitments related to {concern.lower()}."
                    if concern else ""
                ),
                "severity": t.get("severity", "medium"),
            })

    # Recalculate risk level based on validated triggers
    if claude_activist is not None:
        if len(mapped) == 0:
            risk_level = "Low"
        elif any(t.get("severity", "").lower() == "high" for t in mapped):
            risk_level = "High"
        elif len(mapped) >= 3:
            risk_level = "Medium"
        else:
            risk_level = "Low"
    else:
        risk_level = act.get("risk_level", "Low")

    return {
        "risk_level": risk_level,
        "triggers": mapped,
    }


def _build_guidance(
    advanced: dict,
    claude_guidance: Optional[list] = None,
    consensus: Optional[dict] = None,
) -> dict:
    gc = advanced.get("guidance_clarity", {})

    # Flatten signal objects into readable strings for the frontend
    pos_signals = []
    for s in gc.get("positive_signals", []):
        if isinstance(s, dict):
            label = s.get("signal", "")
            count = s.get("count", 0)
            pos_signals.append(f"{label} ({count}x)" if count > 1 else label)
        else:
            pos_signals.append(str(s))

    neg_signals = []
    for s in gc.get("negative_signals", []):
        if isinstance(s, dict):
            label = s.get("signal", "")
            count = s.get("count", 0)
            neg_signals.append(f"{label} ({count}x)" if count > 1 else label)
        else:
            neg_signals.append(str(s))

    # Build metrics list from Claude guidance extraction (preferred) or fallback
    metrics = []
    no_guidance_given = False

    if claude_guidance is not None:
        # Claude extracted actual guidance — use it as the source of truth
        if len(claude_guidance) == 0:
            no_guidance_given = True
        for gm in claude_guidance:
            metric_name = gm.get("metric", "Unknown")
            company_value = gm.get("value", "")
            quantified = gm.get("quantified", False)
            quote = gm.get("quote", "")

            # Try to match against consensus for status
            status = "unknown"
            consensus_value = None
            if consensus:
                consensus_value = _match_consensus(metric_name, consensus)

            if consensus_value is not None and quantified and company_value:
                status = _compare_guidance_to_consensus(company_value, consensus_value, metric_name)

            metrics.append({
                "metric": metric_name,
                "company_guidance": company_value if company_value else ("Directional" if not quantified else "—"),
                "quote": quote,
                "quantified": quantified,
                "consensus": _format_consensus(consensus_value, metric_name) if consensus_value is not None else None,
                "status": status,
            })

        # Recalculate clarity score based on Claude's actual findings
        quantified_count = sum(1 for m in metrics if m.get("quantified"))
        total_metrics = len(metrics)
        if total_metrics == 0:
            clarity_score = 20  # No guidance = low score
        else:
            # Base: 40 for having any guidance, +10 per quantified metric (up to 60 bonus)
            clarity_score = min(100, 40 + (quantified_count * 10))
    else:
        # Fallback: no Claude data available — show honest "analysis unavailable" state
        # Don't show sacred code's false-positive-prone regex metrics
        clarity_score = gc.get("clarity_score", 0)
        no_guidance_given = True  # Signal to frontend to show empty state

    return {
        "clarity_score": clarity_score if claude_guidance is not None else gc.get("clarity_score", 0),
        "grade": _corrected_grade(clarity_score if claude_guidance is not None else gc.get("clarity_score", 0)),
        "metrics": metrics,
        "has_consensus": consensus is not None,
        "no_guidance_given": no_guidance_given,
        "analysis_source": "claude" if claude_guidance is not None else "unavailable",
        "positive_signals": pos_signals,
        "negative_signals": neg_signals,
    }


def _match_consensus(metric_name: str, consensus: dict) -> Optional[float]:
    """Match a guidance metric name to the closest FMP consensus field."""
    name_lower = metric_name.lower()
    if any(k in name_lower for k in ("revenue", "sales", "top line")):
        return consensus.get("revenue")
    if any(k in name_lower for k in ("eps", "earnings per share")):
        return consensus.get("earnings")
    if any(k in name_lower for k in ("net income",)):
        return consensus.get("net_income")
    if any(k in name_lower for k in ("ebitda",)):
        return consensus.get("ebitda")
    return None


def _format_consensus(value: Optional[float], metric_name: str) -> str:
    """Format a consensus value for display."""
    if value is None:
        return "N/A"
    name_lower = metric_name.lower()
    if any(k in name_lower for k in ("eps", "earnings per share")):
        return f"${value:.2f}"
    if value >= 1_000_000_000:
        return f"${value / 1_000_000_000:.1f}B"
    if value >= 1_000_000:
        return f"${value / 1_000_000:.0f}M"
    return f"${value:,.0f}"


def _compare_guidance_to_consensus(guidance_str: str, consensus_val: float, metric_name: str) -> str:
    """Compare guidance value to consensus. Returns status string."""
    # Try to extract a numeric value from the guidance string
    numbers = re.findall(r'[\d,]+\.?\d*', guidance_str.replace(",", ""))
    if not numbers:
        return "unknown"

    parsed = []
    for n in numbers:
        try:
            val = float(n)
            # Handle billions/millions in the original string
            if "billion" in guidance_str.lower() or "b" in guidance_str.lower():
                val *= 1_000_000_000
            elif "million" in guidance_str.lower() or "m" in guidance_str.lower():
                val *= 1_000_000
            parsed.append(val)
        except ValueError:
            continue

    if not parsed:
        return "unknown"

    # If it's a range, use the midpoint
    if len(parsed) >= 2:
        guidance_val = (parsed[0] + parsed[-1]) / 2
    else:
        guidance_val = parsed[0]

    # Normalize: if consensus is in absolute terms and guidance looks like EPS
    if consensus_val == 0:
        return "unknown"

    pct_diff = (guidance_val - consensus_val) / abs(consensus_val)

    if pct_diff > 0.02:
        return "above"
    elif pct_diff > -0.02:
        return "inline"
    elif pct_diff > -0.05:
        return "low_end"
    else:
        return "below"


def _build_consensus_divergence_interps(
    claude_guidance: Optional[list],
    consensus: Optional[dict],
) -> list:
    """
    Compare Claude-extracted guidance against FMP consensus estimates.
    Returns negative interpretation entries for any meaningful divergence.
    """
    if not claude_guidance or not consensus:
        return []

    interps = []
    for gm in claude_guidance:
        metric_name = gm.get("metric", "")
        value_str = gm.get("value", "")
        quantified = gm.get("quantified", False)
        quote = gm.get("quote", "")

        if not quantified or not value_str:
            continue

        consensus_val = _match_consensus(metric_name, consensus)
        if consensus_val is None:
            continue

        status = _compare_guidance_to_consensus(value_str, consensus_val, metric_name)
        consensus_formatted = _format_consensus(consensus_val, metric_name)

        if status == "below":
            interps.append({
                "category": "consensus_divergence",
                "original_text": quote or f"{metric_name} guidance: {value_str}",
                "negative_spin": (
                    f"{metric_name} guidance of {value_str} is meaningfully below "
                    f"Street consensus of {consensus_formatted}. Analysts will view this as "
                    f"a guide-down and press for the drivers behind the shortfall. "
                    f"Expect pointed questions on what changed since the prior quarter."
                ),
                "suggested_rewrite": (
                    f"Provide context for the {metric_name} outlook — explain what headwinds "
                    f"are being incorporated and frame it as prudent conservatism rather than "
                    f"deteriorating fundamentals. Consider adding bridge items that show a path "
                    f"back toward consensus."
                ),
                "severity": "high",
            })
        elif status == "low_end":
            interps.append({
                "category": "consensus_divergence",
                "original_text": quote or f"{metric_name} guidance: {value_str}",
                "negative_spin": (
                    f"{metric_name} guidance of {value_str} is at the low end of "
                    f"Street consensus ({consensus_formatted}). While not a clear miss, "
                    f"analysts will note the lack of upside and may interpret this as "
                    f"sandbagging or underlying weakness."
                ),
                "suggested_rewrite": (
                    f"Acknowledge the conservative posture on {metric_name} and provide "
                    f"specific catalysts or upside drivers that could move the number higher. "
                    f"Frame the range as achievable with clear line of sight."
                ),
                "severity": "medium",
            })

    return interps


def _build_prior_comparison(current_scores: dict, prior_scores: list) -> Optional[dict]:
    """Build prior call comparison data from scored prior transcripts."""
    if not prior_scores:
        return None

    result = {"prior_scores": prior_scores}

    # Delta vs most recent prior quarter
    most_recent = prior_scores[0]
    dims = ["overall", "sentiment", "confidence", "ownership", "clarity", "red_flags"]
    dimension_deltas = {}
    for d in dims:
        cur = current_scores.get(d, 0)
        prev = most_recent["scores"].get(d, 0)
        dimension_deltas[d] = cur - prev

    result["vs_prior"] = {
        "quarter": most_recent["quarter"],
        "overall_delta": dimension_deltas["overall"],
        "dimension_deltas": {k: v for k, v in dimension_deltas.items() if k != "overall"},
    }

    return result


def _build_response(
    text: str,
    base_analysis: dict,
    advanced: dict,
    session_id: str,
    claude_qa: Optional[dict] = None,
    claude_rewrites: Optional[dict] = None,
    claude_analysis: Optional[dict] = None,
    consensus: Optional[dict] = None,
    prior_comparison: Optional[dict] = None,
    claude_bull_bear: Optional[dict] = None,
    ai_status: Optional[dict] = None,
) -> dict:
    """Assemble the spec-compliant JSON response."""
    flagged_passages = base_analysis.get("flagged_passages", [])
    scores = _corrected_scores(base_analysis)

    # Use Claude negative interps if available, fall back to sacred code pattern matches
    if claude_analysis and claude_analysis.get("negative_interpretations"):
        neg_interps = claude_analysis["negative_interpretations"]
    else:
        # Map sacred code format to frontend format
        neg_interps = []
        for item in advanced.get("negative_interpretations", []):
            neg_interps.append({
                "category": item.get("suggestion_type", "awareness"),
                "original_text": item.get("sentence", item.get("matched_text", "")),
                "negative_spin": item.get("interpretation", ""),
                "suggested_rewrite": item.get("rewrite_suggestion"),
                "severity": item.get("severity", "medium"),
            })

    # Extract Claude sub-analyses for downstream builders
    claude_guidance = None
    claude_litigation = None
    claude_activist = None
    if claude_analysis:
        if "guidance_metrics" in claude_analysis:
            claude_guidance = claude_analysis["guidance_metrics"]
        if "litigation_findings" in claude_analysis:
            claude_litigation = claude_analysis["litigation_findings"]
        if "activist_triggers" in claude_analysis:
            claude_activist = claude_analysis["activist_triggers"]

    # Inject consensus divergence findings into negative interpretations
    consensus_interps = _build_consensus_divergence_interps(claude_guidance, consensus)
    if consensus_interps:
        neg_interps = neg_interps + consensus_interps
        logger.info(f"Added {len(consensus_interps)} consensus divergence neg interps")

    return {
        "scores": scores,
        "flagged_issues": _build_flagged_issues(text, base_analysis, claude_rewrites, claude_bull_bear),
        "analyst_qa": claude_qa if claude_qa else _build_analyst_qa_fallback(advanced),
        "negative_interpretations": neg_interps,
        "litigation": _build_litigation(advanced, claude_litigation),
        "activist_triggers": _build_activist(advanced, claude_activist),
        "guidance_clarity": _build_guidance(advanced, claude_guidance, consensus),
        "bull_bear_cases": {
            "bull_cases": claude_bull_bear.get("bull_cases", []),
            "bear_cases": claude_bull_bear.get("bear_cases", []),
            "rewrite_count": len(claude_bull_bear.get("rewrites", [])),
        } if claude_bull_bear else None,
        "prior_comparison": prior_comparison,
        "session_id": session_id,
        "ai_status": ai_status or {"api_configured": False, "degraded_services": [], "degraded": False},
        "_debug": {
            "flagged_passage_count": len(flagged_passages),
            "claude_rewrites_is_none": claude_rewrites is None,
            "claude_rewrites_count": len(claude_rewrites) if claude_rewrites else 0,
            "claude_analysis_available": claude_analysis is not None,
            "claude_analysis_parts": list(claude_analysis.keys()) if claude_analysis else [],
            "claude_guidance_count": len(claude_guidance) if claude_guidance else 0,
            "claude_litigation_count": len(claude_litigation) if claude_litigation else 0,
            "claude_activist_count": len(claude_activist) if claude_activist else 0,
            "claude_bull_bear_available": claude_bull_bear is not None,
            "consensus_available": consensus is not None,
        },
    }


# ---------------------------------------------------------------------------
# Word export with word-level diffs (replaces sacred export_word)
# ---------------------------------------------------------------------------

def _export_word_improved(
    text: str,
    analysis: dict,
    output_path: str,
    claude_rewrites: Optional[dict] = None,
    advanced: Optional[dict] = None,
    claude_analysis: Optional[dict] = None,
    prior_comparison: Optional[dict] = None,
    claude_bull_bear: Optional[dict] = None,
):
    """
    Word export with word-level diffs instead of full-sentence strikethrough.
    Uses Claude rewrites when available for context-aware suggestions.
    Includes litigation risk and activist trigger analyses (Claude-validated when available).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()

    doc.add_heading("Earnings Script — Suggested Revisions", 0)
    doc.add_paragraph()

    # Instructions
    instructions = doc.add_paragraph()
    instructions.add_run("Instructions: ").bold = True
    instructions.add_run(
        "This document shows suggested revisions to improve your earnings script. "
        "Red strikethrough text should be removed or replaced. "
        "Blue underlined text shows proposed additions or replacements. "
        "Accept or reject changes as appropriate for your communication style."
    )
    doc.add_paragraph()

    # Legend
    legend = doc.add_paragraph()
    legend.add_run("Legend: ").bold = True
    strike_run = legend.add_run("Strikethrough")
    strike_run.font.strike = True
    strike_run.font.color.rgb = RGBColor(180, 0, 0)
    legend.add_run(" = Remove  |  ")
    add_run = legend.add_run("Blue Text")
    add_run.font.underline = True
    add_run.font.color.rgb = RGBColor(26, 86, 219)
    legend.add_run(" = Addition")
    doc.add_paragraph()
    doc.add_paragraph("\u2500" * 50)
    doc.add_paragraph()

    # Score Comparison section
    scores = _corrected_scores(analysis)
    _has_comparison = prior_comparison and prior_comparison.get("vs_prior")

    if _has_comparison:
        doc.add_heading("Score Comparison", level=1)

        if prior_comparison and prior_comparison.get("vs_prior"):
            vs = prior_comparison["vs_prior"]
            delta = vs.get("overall_delta", 0)
            sign = "+" if delta > 0 else ""
            quarter = vs.get("quarter", "Prior")
            prior_overall = prior_comparison["prior_scores"][0]["scores"]["overall"] if prior_comparison.get("prior_scores") else "?"

            doc.add_heading(f"vs. Prior Call ({quarter})", level=2)

            from docx.shared import Inches
            prior_table = doc.add_table(rows=7, cols=3)
            prior_table.style = "Light Grid Accent 1"

            headers = ["Dimension", "Current", quarter]
            for i, h in enumerate(headers):
                prior_table.rows[0].cells[i].text = h

            dim_map = [
                ("Overall", "overall"),
                ("Sentiment", "sentiment"),
                ("Confidence", "confidence"),
                ("Ownership", "ownership"),
                ("Clarity", "clarity"),
                ("Red Flags", "red_flags"),
            ]
            for row_idx, (label, key) in enumerate(dim_map, start=1):
                cur_val = scores.get(key, 0)
                prev_val = prior_comparison["prior_scores"][0]["scores"].get(key, 0) if prior_comparison.get("prior_scores") else 0
                d = cur_val - prev_val
                d_str = f" ({'+' if d > 0 else ''}{d})"
                prior_table.rows[row_idx].cells[0].text = label
                prior_table.rows[row_idx].cells[1].text = f"{cur_val}{d_str}"
                prior_table.rows[row_idx].cells[2].text = str(prev_val)

            doc.add_paragraph()

            # Trend line
            if prior_comparison.get("prior_scores") and len(prior_comparison["prior_scores"]) > 1:
                trend_parts = [
                    f"{ps['quarter']}: {ps['scores']['overall']}"
                    for ps in reversed(prior_comparison["prior_scores"])
                ]
                trend_parts.append(f"Current: {scores.get('overall', 0)}")
                trend_p = doc.add_paragraph()
                trend_p.add_run("Trend: ").bold = True
                trend_p.add_run(" \u2192 ".join(trend_parts))
                doc.add_paragraph()

        doc.add_paragraph("\u2500" * 50)
        doc.add_paragraph()

    # Build rewrite lookup from flagged_passages + claude_rewrites.
    # Key insight: use the EXACT sentence strings from flagged_passages
    # (same strings sent to Claude) to look up in claude_rewrites, then
    # store with normalized keys for matching against re-split sentences.
    _passage_rewrites = {}  # norm_key → rewrite
    _flagged_count = 0
    for fp in analysis.get("flagged_passages", []):
        sent = fp["sentence"]
        _flagged_count += 1
        rewrite = None
        if claude_rewrites:
            # Exact match — these keys come from the same flagged_passages list
            rewrite = claude_rewrites.get(sent)
        if not rewrite or rewrite == sent:
            rewrite = _generate_rewrite_fallback(sent, fp["issues"])
        if rewrite and rewrite != sent:
            _passage_rewrites[' '.join(sent.split())] = rewrite

    # Also add any claude_rewrites that aren't from flagged_passages
    # (these come from classify_sentence()-flagged sentences added at analyze time)
    _extra_from_claude = 0
    if claude_rewrites:
        for sent, rewrite in claude_rewrites.items():
            norm_key = ' '.join(sent.split())
            if norm_key not in _passage_rewrites and rewrite and rewrite != sent:
                _passage_rewrites[norm_key] = rewrite
                _extra_from_claude += 1

    # Merge negative-interpretation / litigation / activist rewrites into
    # the same pipeline so they all render as inline track changes.
    # Track annotation source for each rewrite so we can label them.
    _rewrite_annotations = {}  # norm_key → annotation string

    _neg_merged = 0
    _lit_merged = 0
    _act_merged = 0

    # Track all advanced rewrites separately so we can render unmatched
    # ones in dedicated fallback sections at the end of the document.
    _advanced_rewrites = []  # list of (source, orig, rewrite, annotation, item_dict)

    if claude_analysis and "negative_interpretations" in claude_analysis:
        for ni in claude_analysis["negative_interpretations"]:
            orig = ni.get("original_text", "")
            rewrite = ni.get("suggested_rewrite", "")
            if orig and rewrite and rewrite != orig:
                norm_key = ' '.join(orig.split())
                # Only add if not already covered by flagged_passages / claude_rewrites
                if norm_key not in _passage_rewrites:
                    _passage_rewrites[norm_key] = rewrite
                    _neg_merged += 1
                category = ni.get("category", "negative interpretation").replace("_", " ").title()
                severity = ni.get("severity", "medium").upper()
                annotation = f"Neg Interp: {category} [{severity}]"
                _rewrite_annotations[norm_key] = annotation
                _advanced_rewrites.append(("neg_interp", orig, rewrite, annotation, ni))

    if claude_analysis and "litigation_findings" in claude_analysis:
        for f in claude_analysis["litigation_findings"]:
            orig = f.get("original_text", "")
            rewrite = f.get("suggested_rewrite", "")
            if orig and rewrite and rewrite != orig:
                norm_key = ' '.join(orig.split())
                _passage_rewrites[norm_key] = rewrite
                severity = f.get("severity", "medium").upper()
                issue = f.get("issue", "Litigation Risk")
                annotation = f"Litigation: {issue} [{severity}]"
                _rewrite_annotations[norm_key] = annotation
                _lit_merged += 1
                _advanced_rewrites.append(("litigation", orig, rewrite, annotation, f))

    if claude_analysis and "activist_triggers" in claude_analysis:
        for t in claude_analysis["activist_triggers"]:
            orig = t.get("original_text", "")
            rewrite = t.get("suggested_rewrite", "")
            if orig and rewrite and rewrite != orig:
                norm_key = ' '.join(orig.split())
                # Don't overwrite a litigation rewrite (higher priority)
                if norm_key not in _passage_rewrites:
                    _passage_rewrites[norm_key] = rewrite
                if norm_key not in _rewrite_annotations:
                    severity = t.get("severity", "medium").upper()
                    category = t.get("category", "Activist Vulnerability")
                    annotation = f"Activist: {category} [{severity}]"
                    _rewrite_annotations[norm_key] = annotation
                else:
                    annotation = _rewrite_annotations[norm_key]
                _act_merged += 1
                _advanced_rewrites.append(("activist", orig, rewrite, annotation, t))

    # Bull/Bear Defense rewrites — these OVERWRITE existing rewrites for the
    # same sentence because a well-crafted bull/bear rewrite addresses both the
    # market narrative AND the underlying issue (litigation, activist, etc.)
    _bb_merged = 0
    if claude_bull_bear and "rewrites" in claude_bull_bear:
        for rw in claude_bull_bear["rewrites"]:
            if not isinstance(rw, dict):
                continue
            orig = rw.get("original_text", "")
            rewrite = rw.get("suggested_rewrite", "")
            if orig and rewrite and rewrite != orig:
                norm_key = ' '.join(orig.split())
                # Note if we're overwriting an existing rewrite
                existing_annotation = _rewrite_annotations.get(norm_key)
                also_note = f" [also: {existing_annotation}]" if existing_annotation else ""
                # Overwrite — bull/bear takes priority
                _passage_rewrites[norm_key] = rewrite
                case_type = rw.get("case_type", "bear").title()
                thesis_full = rw.get("case_thesis", "")
                if len(thesis_full) > 80:
                    # Truncate at last word boundary before 80 chars, add ellipsis
                    thesis_short = thesis_full[:80].rsplit(" ", 1)[0] + "…"
                else:
                    thesis_short = thesis_full
                annotation = f"Bull/Bear: {case_type} — {thesis_short}{also_note}"
                _rewrite_annotations[norm_key] = annotation
                _bb_merged += 1
                _advanced_rewrites.append(("bull_bear", orig, rewrite, annotation, rw))

    logger.info(f"Word export: {_flagged_count} flagged passages, "
                 f"{len(_passage_rewrites)} with rewrites "
                 f"({_extra_from_claude} extra from classify_sentence via claude_rewrites, "
                 f"{_neg_merged} neg interps, {_lit_merged} litigation, {_act_merged} activist, "
                 f"{_bb_merged} bull/bear), "
                 f"claude_rewrites={'None' if claude_rewrites is None else len(claude_rewrites)}")

    # Process sentences
    # Split on sentence-ending punctuation, then re-join fragments that
    # were incorrectly split on abbreviation periods (Inc., Corp., etc.)
    _ABBREVS = {"Inc", "Corp", "Ltd", "Co", "Dr", "Mr", "Mrs", "Ms",
                "Jr", "Sr", "St", "vs", "etc", "Vol", "No", "Dept",
                "Prof", "Gen", "Gov", "Rev", "Hon", "Pres", "Assn"}
    raw_parts = re.split(r"(?<=[.!?])\s+", text)
    sentences = []
    for part in raw_parts:
        # If previous fragment ends with an abbreviation period, merge
        if sentences and any(sentences[-1].rstrip(".").endswith(ab) for ab in _ABBREVS):
            sentences[-1] = sentences[-1] + " " + part
        else:
            sentences.append(part)
    _exact_hits = 0
    _fuzzy_hits = 0
    _word_overlap_hits = 0
    _matched_advanced_keys = set()  # Track which advanced rewrite keys matched inline
    _inline_rendered = {}  # norm_key -> [(actual_sentence, actual_rewrite)] for summary consistency

    for sentence in sentences:
        if len(sentence.strip()) < 10:
            continue

        color, issues = classify_sentence(sentence)

        # Step 1: Try normalized exact match
        norm_sentence = ' '.join(sentence.split())
        rewrite = _passage_rewrites.get(norm_sentence)
        matched_key = norm_sentence if rewrite else None

        if rewrite:
            _exact_hits += 1
        elif _passage_rewrites and (color in ("RED", "YELLOW") or _rewrite_annotations):
            # Step 2: Fuzzy match — sentence boundaries from regex splitter
            # may differ from sacred code's splitter.
            # Also runs for GREEN/NEUTRAL sentences when litigation/activist
            # rewrites exist, since those quotes may not be flagged by
            # classify_sentence().
            best_ratio = 0
            best_rw = None
            best_key = None
            for key, rw in _passage_rewrites.items():
                # Step 2a: Check substring containment — Claude's original_text
                # may span multiple sentences. If the current sentence is
                # contained within a multi-sentence key, it's a match.
                # Guard: the sentence must be at least 40% of the key length
                # to prevent tiny fragments (e.g. "Owlet, Inc.") from matching
                # against long multi-sentence passages.
                if (len(key) > len(norm_sentence) + 10
                        and norm_sentence in key
                        and len(norm_sentence) >= len(key) * 0.4):
                    # Sentence is part of a multi-sentence original_text
                    ratio = 0.95
                elif (len(norm_sentence) > len(key) + 10
                        and key in norm_sentence
                        and len(key) >= len(norm_sentence) * 0.4):
                    # Key (Claude's original_text) is a substring of the
                    # actual script sentence — Claude quoted only part of it.
                    # Build a full-sentence rewrite by splicing the rewrite
                    # into the sentence at the matched position.
                    idx = norm_sentence.index(key)
                    prefix = norm_sentence[:idx]
                    suffix = norm_sentence[idx + len(key):]
                    spliced_rw = prefix + rw + suffix
                    # Store the spliced version for this match
                    if spliced_rw != norm_sentence:
                        ratio = 0.95
                        rw = spliced_rw  # local override for this candidate
                    else:
                        continue
                elif abs(len(key) - len(norm_sentence)) > 30:
                    continue
                else:
                    ratio = difflib.SequenceMatcher(None, norm_sentence, key).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_rw = rw
                    best_key = key
            if best_ratio > 0.8:
                rewrite = best_rw
                matched_key = best_key
                _fuzzy_hits += 1

        # Step 3: Word-overlap match against advanced rewrite originals.
        # Handles character encoding differences (smart quotes, em-dashes)
        # and cases where the fuzzy length threshold (30 chars) skips the key.
        _STOPWORDS = {
            "a", "an", "the", "and", "or", "but", "in", "on", "at", "to",
            "for", "of", "with", "by", "from", "is", "are", "was", "were",
            "be", "been", "being", "have", "has", "had", "do", "does", "did",
            "will", "would", "shall", "should", "may", "might", "must", "can",
            "could", "not", "no", "so", "if", "as", "it", "its", "we", "our",
            "us", "this", "that", "these", "those", "than", "into", "also",
        }
        if (not rewrite or rewrite == sentence) and _advanced_rewrites:
            sent_words = norm_sentence.lower().split()
            sent_content = [w for w in sent_words if w not in _STOPWORDS]
            if len(sent_content) >= 3:
                best_adv_ratio = 0
                best_adv_rw = None
                best_adv_key = None
                for _src, adv_orig, adv_rw, _ann, _item in _advanced_rewrites:
                    adv_norm = ' '.join(adv_orig.split())
                    adv_content = [w for w in adv_norm.lower().split() if w not in _STOPWORDS]
                    adv_content_set = set(adv_content)
                    # Sentence must be ≥70% of original word count to prevent
                    # matching a single sentence against a multi-sentence original.
                    if len(sent_content) < len(adv_content) * 0.7:
                        continue
                    # Forward overlap: what % of sentence content words are in the original
                    fwd_matching = sum(1 for w in sent_content if w in adv_content_set)
                    fwd_overlap = fwd_matching / len(sent_content)
                    # Reverse overlap: what % of original content words are in the sentence
                    sent_content_set = set(sent_content)
                    rev_matching = sum(1 for w in adv_content if w in sent_content_set)
                    rev_overlap = rev_matching / len(adv_content) if adv_content else 0
                    # When the sentence is significantly longer than the original,
                    # Claude likely quoted only part of the sentence — use reverse
                    # overlap (how much of the original is in the sentence) as the
                    # primary signal.  Otherwise use min of both directions.
                    if len(sent_content) > len(adv_content) * 1.3 and rev_overlap > 0.85:
                        overlap = rev_overlap
                    else:
                        overlap = min(fwd_overlap, rev_overlap)
                    if overlap > best_adv_ratio:
                        best_adv_ratio = overlap
                        best_adv_rw = adv_rw
                        best_adv_key = adv_norm
                if best_adv_ratio > 0.8:
                    # If key is a substring of the sentence, splice the
                    # rewrite into the full sentence to avoid deleting the
                    # unmatched tail/head.
                    if (best_adv_key
                            and len(norm_sentence) > len(best_adv_key) + 10
                            and best_adv_key in norm_sentence):
                        idx = norm_sentence.index(best_adv_key)
                        prefix = norm_sentence[:idx]
                        suffix = norm_sentence[idx + len(best_adv_key):]
                        best_adv_rw = prefix + best_adv_rw + suffix
                    rewrite = best_adv_rw
                    matched_key = best_adv_key
                    _word_overlap_hits += 1
                    # Ensure key is in _passage_rewrites for tracking
                    if best_adv_key not in _passage_rewrites:
                        _passage_rewrites[best_adv_key] = best_adv_rw

        if not rewrite or rewrite == sentence:
            rewrite = _generate_rewrite_fallback(sentence, issues)

        para = doc.add_paragraph()

        # Look up litigation/activist annotation for this sentence
        annotation = None
        if matched_key:
            annotation = _rewrite_annotations.get(matched_key)
            if annotation:
                _matched_advanced_keys.add(matched_key)
                # Track the actual sentence + rewrite used inline so summary
                # sections render the exact same diff as the script body.
                if rewrite and rewrite != sentence:
                    _inline_rendered.setdefault(matched_key, []).append((sentence, rewrite))

        if rewrite and rewrite != sentence:
            # Word-level diff
            orig_words = sentence.split()
            new_words = rewrite.split()
            matcher = difflib.SequenceMatcher(None, orig_words, new_words)

            for op, i1, i2, j1, j2 in matcher.get_opcodes():
                if op == "equal":
                    para.add_run(" ".join(orig_words[i1:i2]) + " ")
                elif op == "replace":
                    # Strikethrough the old words
                    old_run = para.add_run(" ".join(orig_words[i1:i2]) + " ")
                    old_run.font.strike = True
                    old_run.font.color.rgb = RGBColor(180, 0, 0)
                    # Blue underline the new words
                    new_run = para.add_run(" ".join(new_words[j1:j2]) + " ")
                    new_run.font.underline = True
                    new_run.font.color.rgb = RGBColor(26, 86, 219)
                elif op == "delete":
                    old_run = para.add_run(" ".join(orig_words[i1:i2]) + " ")
                    old_run.font.strike = True
                    old_run.font.color.rgb = RGBColor(180, 0, 0)
                elif op == "insert":
                    new_run = para.add_run(" ".join(new_words[j1:j2]) + " ")
                    new_run.font.underline = True
                    new_run.font.color.rgb = RGBColor(26, 86, 219)

            # Build annotation label from issues + litigation/activist source
            labels = []
            if issues:
                labels.extend(issues)
            if annotation:
                labels.append(annotation)
            if labels:
                comment = para.add_run("  [" + ", ".join(labels) + "]")
                comment.font.size = Pt(8)
                comment.font.italic = True
                comment.font.color.rgb = RGBColor(128, 128, 128)

        elif color in ("RED", "YELLOW") and issues:
            # No rewrite available but flagged — show with highlight + issue note
            run = para.add_run(sentence)
            if color == "RED":
                run.font.highlight_color = WD_COLOR_INDEX.PINK
            else:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            comment = para.add_run("  [" + ", ".join(issues) + "]")
            comment.font.size = Pt(8)
            comment.font.italic = True
            comment.font.color.rgb = RGBColor(128, 128, 128)
        else:
            para.add_run(sentence)

        para.add_run(" ")

    logger.info(f"Word export complete: {_exact_hits} exact + {_fuzzy_hits} fuzzy + "
                 f"{_word_overlap_hits} word-overlap matches out of {len(sentences)} sentences, "
                 f"{len(_matched_advanced_keys)} advanced rewrites matched inline")

    # Helper: render a single rewrite item with track-changes formatting
    def _render_rewrite_item(doc, orig_text, rewrite_text, label):
        """Render original → rewrite with word-level diffs in the document."""
        para = doc.add_paragraph()
        # Label
        lbl_run = para.add_run(f"[{label}]  ")
        lbl_run.font.size = Pt(9)
        lbl_run.font.bold = True
        lbl_run.font.color.rgb = RGBColor(128, 128, 128)

        # Word-level diff
        orig_words = orig_text.split()
        new_words = rewrite_text.split()
        matcher = difflib.SequenceMatcher(None, orig_words, new_words)

        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == "equal":
                para.add_run(" ".join(orig_words[i1:i2]) + " ")
            elif op == "replace":
                old_run = para.add_run(" ".join(orig_words[i1:i2]) + " ")
                old_run.font.strike = True
                old_run.font.color.rgb = RGBColor(180, 0, 0)
                new_run = para.add_run(" ".join(new_words[j1:j2]) + " ")
                new_run.font.underline = True
                new_run.font.color.rgb = RGBColor(26, 86, 219)
            elif op == "delete":
                old_run = para.add_run(" ".join(orig_words[i1:i2]) + " ")
                old_run.font.strike = True
                old_run.font.color.rgb = RGBColor(180, 0, 0)
            elif op == "insert":
                new_run = para.add_run(" ".join(new_words[j1:j2]) + " ")
                new_run.font.underline = True
                new_run.font.color.rgb = RGBColor(26, 86, 219)
        doc.add_paragraph()  # spacer

    # Collect unmatched advanced rewrites by source
    _unmatched_neg = []
    _unmatched_lit = []
    _unmatched_act = []
    for source, orig, rewrite, annotation, item in _advanced_rewrites:
        norm_key = ' '.join(orig.split())
        if norm_key not in _matched_advanced_keys:
            if source == "neg_interp":
                _unmatched_neg.append((orig, rewrite, annotation, item))
            elif source == "litigation":
                _unmatched_lit.append((orig, rewrite, annotation, item))
            elif source == "activist":
                _unmatched_act.append((orig, rewrite, annotation, item))

    # --- Negative Interpretation Rewrites ---
    # (These rewrites are rendered inline in the script body above.
    #  No separate section needed — see _passage_rewrites merge at ~line 1557.)

    # --- Litigation Risk Analysis Summary ---
    if advanced:
        lit = advanced.get("litigation_risk", {})
        has_safe_harbor = lit.get("has_safe_harbor", False)

        if claude_analysis and "litigation_findings" in claude_analysis:
            findings = claude_analysis["litigation_findings"]
        else:
            findings = lit.get("findings", [])

        if findings or not has_safe_harbor:
            doc.add_paragraph()
            doc.add_paragraph("\u2500" * 50)
            doc.add_heading("Litigation Risk Summary", level=1)

            # Safe harbor status
            sh_para = doc.add_paragraph()
            if has_safe_harbor:
                sh_run = sh_para.add_run("\u2713 Safe Harbor Statement Present")
                sh_run.font.color.rgb = RGBColor(5, 150, 105)
                sh_run.bold = True
            else:
                sh_run = sh_para.add_run("\u2717 No Safe Harbor Statement Detected")
                sh_run.font.color.rgb = RGBColor(220, 38, 38)
                sh_run.bold = True
                rec = doc.add_paragraph()
                rec.add_run(
                    "Recommendation: Add a PSLRA-compliant forward-looking statement disclaimer "
                    "at the beginning of the prepared remarks."
                ).font.italic = True

            # Risk level
            if findings:
                if any(f.get("severity", "").lower() in ("high", "critical") for f in findings):
                    risk_level = "High"
                elif len(findings) >= 3:
                    risk_level = "Medium"
                else:
                    risk_level = "Low"
            else:
                risk_level = "Low"
            doc.add_paragraph(f"Overall Risk: {risk_level}  ({len(findings)} finding{'s' if len(findings) != 1 else ''})")
            doc.add_paragraph()

            # Show each finding with its rewrite
            for f in findings:
                severity = f.get("severity", "medium").upper()
                issue = f.get("issue", "Finding")
                orig = f.get("original_text", "")
                rewrite = f.get("suggested_rewrite", "")
                label = f"Litigation: {issue} [{severity}]"

                if orig and rewrite and rewrite != orig:
                    norm_key = ' '.join(orig.split())
                    if norm_key in _inline_rendered:
                        for actual_sentence, actual_rewrite in _inline_rendered[norm_key]:
                            _render_rewrite_item(doc, actual_sentence, actual_rewrite, label)
                    else:
                        _render_rewrite_item(doc, orig, rewrite, label)
                else:
                    # No rewrite — just show the finding as a bullet
                    bullet = doc.add_paragraph(style="List Bullet")
                    sev_run = bullet.add_run(f"[{severity}] ")
                    sev_run.font.size = Pt(9)
                    if severity in ("CRITICAL", "HIGH"):
                        sev_run.font.color.rgb = RGBColor(220, 38, 38)
                    elif severity == "MEDIUM":
                        sev_run.font.color.rgb = RGBColor(217, 119, 6)
                    else:
                        sev_run.font.color.rgb = RGBColor(107, 114, 128)
                    bullet.add_run(issue)
                    if orig:
                        quote_para = doc.add_paragraph()
                        quote_run = quote_para.add_run(f'"{orig}"')
                        quote_run.font.italic = True
                        quote_run.font.size = Pt(9)

        # --- Activist Vulnerability Summary ---
        # (These rewrites are rendered inline in the script body above.
        #  No separate section needed — see _passage_rewrites merge at ~line 1587.)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Waitlist
# ---------------------------------------------------------------------------
WAITLIST_FILE = os.path.join(os.path.dirname(__file__), "waitlist.json")

@app.post("/api/waitlist")
async def join_waitlist(request: Request):
    """Add an email to the early-access waitlist."""
    body = await request.json()
    email = body.get("email", "").strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Invalid email")

    # Load existing waitlist
    entries = []
    if os.path.exists(WAITLIST_FILE):
        with open(WAITLIST_FILE, "r") as f:
            entries = json.load(f)

    # Avoid duplicates
    if any(e["email"] == email for e in entries):
        return {"status": "already_registered"}

    entries.append({"email": email, "joined_at": datetime.now().isoformat()})
    with open(WAITLIST_FILE, "w") as f:
        json.dump(entries, f, indent=2)

    logger.info(f"Waitlist signup: {email}")
    return {"status": "ok"}


@app.get("/api/health")
async def health():
    """Diagnostic endpoint — check if API keys are configured."""
    return {
        "status": "ok",
        "anthropic_key_set": bool(ANTHROPIC_API_KEY),
        "anthropic_key_prefix": ANTHROPIC_API_KEY[:12] + "..." if ANTHROPIC_API_KEY else "NOT SET",
        "fmp_key_set": bool(FMP_API_KEY),
    }


@app.get("/api/test-fmp/{ticker}")
async def test_fmp(ticker: str):
    """Diagnostic: test FMP API connectivity for a ticker."""
    if not FMP_API_KEY:
        return {"status": "error", "detail": "FMP_API_KEY not set"}

    ticker = ticker.strip().upper()
    results = {}

    # Test transcript endpoint
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.get(
                "https://financialmodelingprep.com/stable/earning-call-transcript",
                params={"symbol": ticker, "quarter": 4, "year": 2025, "apikey": FMP_API_KEY},
            )
            results["transcript_status"] = resp.status_code
            results["transcript_response_length"] = len(resp.text)
            data = resp.json()
            if isinstance(data, list) and len(data) > 0:
                results["transcript_found"] = True
                results["transcript_content_length"] = len(data[0].get("content", ""))
            elif isinstance(data, dict) and "Error Message" in str(data):
                results["transcript_found"] = False
                results["transcript_error"] = str(data)[:200]
            else:
                results["transcript_found"] = False
                results["transcript_raw"] = str(data)[:200]
    except Exception as e:
        results["transcript_error"] = f"{type(e).__name__}: {e}"

    # Test peers endpoint
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(
                "https://financialmodelingprep.com/stable/stock-peers",
                params={"symbol": ticker, "apikey": FMP_API_KEY},
            )
            results["peers_status"] = resp.status_code
            data = resp.json()
            if isinstance(data, list) and len(data) > 0:
                results["peers_found"] = data[0].get("peersList", [])[:5]
            elif isinstance(data, dict):
                results["peers_error"] = str(data)[:200]
            else:
                results["peers_raw"] = str(data)[:200]
    except Exception as e:
        results["peers_error"] = f"{type(e).__name__}: {e}"

    return {"status": "ok", "ticker": ticker, **results}


@app.get("/api/test-claude")
async def test_claude():
    """Test Claude API connectivity — makes a tiny call to verify the key works."""
    if not ANTHROPIC_API_KEY:
        return {"status": "error", "detail": "ANTHROPIC_API_KEY not set"}
    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        response = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=50,
            messages=[{"role": "user", "content": "Say hello in exactly 5 words."}],
        )
        return {
            "status": "ok",
            "model": response.model,
            "response": response.content[0].text,
        }
    except Exception as e:
        return {"status": "error", "error_type": type(e).__name__, "detail": str(e)}


@app.get("/api/debug-session/{session_id}")
async def debug_session(session_id: str):
    """Diagnostic: shows rewrite pipeline status for a session."""
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    claude_rewrites = session.get("claude_rewrites")
    flagged = session["base_analysis"].get("flagged_passages", [])

    # Check how many flagged passages would get rewrites
    match_count = 0
    if claude_rewrites and flagged:
        for fp in flagged:
            sent = fp["sentence"]
            if sent in claude_rewrites or ' '.join(sent.split()) in claude_rewrites:
                match_count += 1

    # Count how many classify_sentence would flag beyond flagged_passages
    text = session.get("text", "")
    classify_extra = 0
    flagged_norm = {' '.join(fp["sentence"].split()) for fp in flagged}
    if text:
        for sent in re.split(r"(?<=[.!?])\s+", text):
            if len(sent.strip()) < 10:
                continue
            norm = ' '.join(sent.split())
            if norm in flagged_norm:
                continue
            color, _ = classify_sentence(sent)
            if color in ("RED", "YELLOW"):
                classify_extra += 1

    return {
        "flagged_passage_count": len(flagged),
        "classify_sentence_extra": classify_extra,
        "total_sent_to_claude": len(flagged) + classify_extra,
        "flagged_sample": [fp["sentence"][:100] for fp in flagged[:5]],
        "claude_rewrites_is_none": claude_rewrites is None,
        "claude_rewrites_count": len(claude_rewrites) if claude_rewrites else 0,
        "claude_rewrite_keys_sample": [k[:100] for k in list(claude_rewrites.keys())[:5]] if claude_rewrites else [],
        "matched_rewrites": match_count,
        "claude_qa_source": session.get("claude_qa", {}).get("source", "fallback") if session.get("claude_qa") else "fallback",
    }


# ---------------------------------------------------------------------------
# Async analyze flow: start → poll status → fetch result
# Lets the UI show real per-Claude-call progress instead of a single blocking
# spinner for 45-60s.
# ---------------------------------------------------------------------------

# Hold a strong reference to background tasks so Python's GC doesn't reclaim
# them mid-execution. Per Python docs: asyncio.create_task() only returns a
# weak reference; without this set the task can silently disappear.
_BACKGROUND_TASKS: set = set()


def _spawn_background(coro):
    task = asyncio.create_task(coro)
    _BACKGROUND_TASKS.add(task)
    task.add_done_callback(_BACKGROUND_TASKS.discard)
    return task


async def _extract_transcript(file: Optional[UploadFile], text: Optional[str]) -> str:
    """Pull transcript text from a file upload or form text field."""
    if file is not None:
        raw = await file.read()
        fname = (file.filename or "").lower()
        if fname.endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        if fname.endswith(".pdf"):
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        return raw.decode("utf-8", errors="replace")
    return text or ""


async def _run_claude_analyses_background(
    session_id: str,
    transcript: str,
    ticker: Optional[str],
    base_analysis: dict,
    advanced: dict,
):
    """
    Run all non-instant work (prior-transcript fetch, consensus fetch, 4 Claude
    calls, prior_comparison scoring) in the background. Updates
    SESSIONS[session_id]["progress"] as each Claude call resolves so the
    frontend can show real per-call progress.
    """
    session = SESSIONS[session_id]
    logger.info(f"Background task started for session {session_id} (ticker={ticker})")
    try:
        # ----- Fetch prior transcripts + consensus (ticker-gated) -----
        prior_transcripts: list = []
        consensus = None
        if ticker and ticker.strip():
            try:
                prior_transcripts = await _fetch_prior_transcripts(ticker.strip())
                logger.info(
                    f"Fetched {len(prior_transcripts)} prior transcripts for {ticker.strip().upper()}"
                )
            except Exception as e:
                logger.error(f"Prior transcript fetch failed: {type(e).__name__}: {e}")
            try:
                consensus = await _fetch_consensus_estimates(ticker.strip())
                if consensus:
                    logger.info(f"FMP consensus available for {ticker}: date={consensus.get('date')}")
            except Exception as e:
                logger.warning(f"Consensus fetch failed: {type(e).__name__}: {e}")
        session["consensus"] = consensus

        # ----- Build rewrite candidate list (same logic as the sync endpoint) -----
        flagged_for_rewrite: list = []
        if ANTHROPIC_API_KEY:
            flagged_for_rewrite = [
                {"sentence": fp["sentence"], "issues": fp["issues"]}
                for fp in base_analysis.get("flagged_passages", [])
            ]
            existing_norm = {' '.join(fp["sentence"].split()) for fp in flagged_for_rewrite}
            all_sentences = re.split(r"(?<=[.!?])\s+", transcript)
            for sent in all_sentences:
                if len(sent.strip()) < 10:
                    continue
                norm_sent = ' '.join(sent.split())
                if norm_sent in existing_norm:
                    continue
                color, issues = classify_sentence(sent)
                if color in ("RED", "YELLOW") and issues:
                    flagged_for_rewrite.append({"sentence": sent, "issues": issues})
                    existing_norm.add(norm_sent)
            if len(flagged_for_rewrite) > 80:
                logger.warning(f"Capping rewrite batch from {len(flagged_for_rewrite)} to 80 sentences")
                flagged_for_rewrite = flagged_for_rewrite[:80]

        # ----- Kick off the 4 Claude calls, each tracking its own progress -----
        claude_qa = None
        claude_rewrites = None
        claude_analysis = None
        claude_bull_bear = None

        async def _track(name: str, coro):
            """Run a Claude call, updating session['progress'][name] as it resolves."""
            session["progress"][name] = "running"
            t0 = time.monotonic()
            try:
                result = await coro
                elapsed_ms = int((time.monotonic() - t0) * 1000)
                session["progress"][name] = "complete" if result is not None else "failed"
                session["timings"][name] = elapsed_ms
                return result
            except Exception as e:
                elapsed_ms = int((time.monotonic() - t0) * 1000)
                logger.error(f"Claude {name} failed: {type(e).__name__}: {e}")
                session["progress"][name] = "failed"
                session["timings"][name] = elapsed_ms
                return None

        if ANTHROPIC_API_KEY:
            logger.info(f"Background: starting Claude calls (ticker={ticker})")

            # Build the task list dynamically so we can mark "no work to do"
            # services as "skipped" instead of running them and getting None,
            # which the tracker would misinterpret as a failure.
            tasks = []
            task_names = []

            tasks.append(_track("qa", _generate_qa_with_claude(transcript, prior_transcripts, base_analysis)))
            task_names.append("qa")

            if flagged_for_rewrite:
                tasks.append(_track("rewrites", _generate_rewrites_with_claude(flagged_for_rewrite)))
                task_names.append("rewrites")
            else:
                logger.info("Skipping rewrites: no flagged sentences in transcript")
                session["progress"]["rewrites"] = "skipped"

            tasks.append(_track("analysis", _generate_analysis_with_claude(
                transcript,
                advanced.get("negative_interpretations", []),
                advanced.get("litigation_risk", {}).get("findings", []),
                advanced.get("activist_triggers", {}).get("triggers", []),
            )))
            task_names.append("analysis")

            if ticker and ticker.strip():
                tasks.append(_track("bull_bear", _generate_bull_bear_with_claude(transcript, ticker.strip().upper())))
                task_names.append("bull_bear")
            # else: bull_bear was already initialized to "skipped" at session-create time

            results = await asyncio.gather(*tasks)
            for name, result in zip(task_names, results):
                if name == "qa":
                    claude_qa = result
                elif name == "rewrites":
                    claude_rewrites = result
                elif name == "analysis":
                    claude_analysis = result
                elif name == "bull_bear":
                    claude_bull_bear = result
        else:
            logger.info("No ANTHROPIC_API_KEY — using fallback Q&A and rewrites")

        # ----- Build ai_status (degradation summary for the UI banner) -----
        _has_ticker = bool(ticker and ticker.strip())
        _expected = {
            "qa": bool(ANTHROPIC_API_KEY),
            "rewrites": bool(ANTHROPIC_API_KEY) and len(flagged_for_rewrite) > 0,
            "analysis": bool(ANTHROPIC_API_KEY),
            "bull_bear": bool(ANTHROPIC_API_KEY) and _has_ticker,
        }
        _succeeded = {
            "qa": claude_qa is not None,
            "rewrites": claude_rewrites is not None,
            "analysis": claude_analysis is not None,
            "bull_bear": claude_bull_bear is not None,
        }
        _degraded = [n for n, exp in _expected.items() if exp and not _succeeded[n]]
        ai_status = {
            "api_configured": bool(ANTHROPIC_API_KEY),
            "degraded_services": _degraded,
            "degraded": bool(ANTHROPIC_API_KEY) and len(_degraded) > 0,
        }
        if ai_status["degraded"]:
            logger.warning(f"AI degraded — falling back to templates for: {_degraded}")

        # ----- Score prior transcripts for comparison -----
        prior_scored = []
        if prior_transcripts:
            for t in prior_transcripts:
                try:
                    prior_base = analyze_transcript(t["content"], LM_DICT)
                    prior_sc = _corrected_scores(prior_base)
                    prior_scored.append({"quarter": t["quarter"], "scores": prior_sc})
                except Exception as e:
                    logger.warning(f"Failed to score prior transcript {t['quarter']}: {e}")
        prior_comparison = _build_prior_comparison(_corrected_scores(base_analysis), prior_scored)

        # ----- Persist the full result into the session -----
        session["claude_qa"] = claude_qa
        session["claude_rewrites"] = claude_rewrites
        session["claude_analysis"] = claude_analysis
        session["claude_bull_bear"] = claude_bull_bear
        session["prior_comparison"] = prior_comparison
        session["ai_status"] = ai_status
        session["status"] = "complete"
        session["completed_at"] = time.monotonic()
        logger.info(
            f"Background analysis complete for {session_id} "
            f"(total {int((session['completed_at'] - session['started_monotonic']) * 1000)}ms)"
        )
    except Exception as e:
        logger.error(f"Background analysis failed for {session_id}: {type(e).__name__}: {e}")
        session["status"] = "failed"
        session["error"] = f"{type(e).__name__}: {e}"


@app.post("/api/analyze/start")
async def analyze_start(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    ticker: Optional[str] = Form(None),
):
    """
    Start an analysis in the background. Returns a session_id immediately so
    the UI can begin polling /api/analyze/status/{session_id} for per-call
    progress, then fetch the full result via /api/analyze/result/{session_id}
    when status == 'complete'.
    """
    transcript = await _extract_transcript(file, text)
    if not transcript or len(transcript.strip()) < 100:
        raise HTTPException(
            status_code=400,
            detail="Please provide a transcript of at least 100 characters.",
        )

    # Fast synchronous base analysis — usually <500ms, fine to block on
    base_analysis = analyze_transcript(transcript, LM_DICT)
    advanced = run_advanced_analysis(transcript, base_analysis)

    session_id = str(uuid.uuid4())
    _has_ticker = bool(ticker and ticker.strip())
    SESSIONS[session_id] = {
        "text": transcript,
        "base_analysis": base_analysis,
        "advanced": advanced,
        "ticker": ticker.strip().upper() if _has_ticker else None,
        "status": "running",
        "started_monotonic": time.monotonic(),
        "progress": {
            "qa": "pending" if ANTHROPIC_API_KEY else "skipped",
            "rewrites": "pending" if ANTHROPIC_API_KEY else "skipped",
            "analysis": "pending" if ANTHROPIC_API_KEY else "skipped",
            "bull_bear": "pending" if (ANTHROPIC_API_KEY and _has_ticker) else "skipped",
        },
        "timings": {},
        "claude_qa": None,
        "claude_rewrites": None,
        "claude_analysis": None,
        "claude_bull_bear": None,
        "consensus": None,
        "prior_comparison": None,
        "ai_status": None,
    }

    _spawn_background(
        _run_claude_analyses_background(session_id, transcript, ticker, base_analysis, advanced)
    )

    return {"session_id": session_id, "status": "running"}


@app.get("/api/analyze/status/{session_id}")
async def analyze_status(session_id: str):
    """Return current progress for each Claude call + overall status."""
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    elapsed_ms = int((time.monotonic() - session.get("started_monotonic", time.monotonic())) * 1000)
    return {
        "session_id": session_id,
        "status": session.get("status", "running"),
        "progress": session.get("progress", {}),
        "timings": session.get("timings", {}),
        "elapsed_ms": elapsed_ms,
        "error": session.get("error"),
    }


@app.get("/api/analyze/result/{session_id}")
async def analyze_result(session_id: str):
    """Return the full analysis result once status == 'complete'."""
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.get("status") == "failed":
        raise HTTPException(status_code=500, detail=session.get("error") or "Analysis failed")
    if session.get("status") != "complete":
        raise HTTPException(status_code=409, detail="Analysis still running")

    try:
        return _build_response(
            session["text"],
            session["base_analysis"],
            session["advanced"],
            session_id,
            session.get("claude_qa"),
            session.get("claude_rewrites"),
            session.get("claude_analysis"),
            session.get("consensus"),
            session.get("prior_comparison"),
            session.get("claude_bull_bear"),
            session.get("ai_status"),
        )
    except Exception as e:
        logger.exception(f"_build_response failed for session {session_id}")
        raise HTTPException(status_code=500, detail=f"Result assembly failed: {type(e).__name__}: {e}")


@app.post("/api/analyze")
async def analyze(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    ticker: Optional[str] = Form(None),
):
    # Extract transcript text from either source
    transcript = None
    if file is not None:
        raw = await file.read()
        fname = (file.filename or "").lower()
        if fname.endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            transcript = "\n".join(p.text for p in doc.paragraphs)
        elif fname.endswith(".pdf"):
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            transcript = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        else:
            transcript = raw.decode("utf-8", errors="replace")
    elif text:
        transcript = text

    if not transcript or len(transcript.strip()) < 100:
        raise HTTPException(
            status_code=400,
            detail="Please provide a transcript of at least 100 characters.",
        )

    # Run base analysis (earnings_analyzer) — synchronous, fast
    base_analysis = analyze_transcript(transcript, LM_DICT)

    # Run all advanced analyses via the wrapper — synchronous, fast
    advanced = run_advanced_analysis(transcript, base_analysis)

    # Fetch prior transcripts (FMP only — no Claude needed) and consensus estimates
    prior_transcripts = []
    consensus = None
    if ticker and ticker.strip():
        try:
            prior_transcripts = await _fetch_prior_transcripts(ticker.strip())
            logger.info(
                f"Fetched {len(prior_transcripts)} prior transcripts for {ticker.strip().upper()}"
            )
        except Exception as e:
            logger.error(f"Prior transcript fetch failed: {type(e).__name__}: {e}")

        try:
            consensus = await _fetch_consensus_estimates(ticker.strip())
            if consensus:
                logger.info(f"FMP consensus available for {ticker}: date={consensus.get('date')}")
        except Exception as e:
            logger.warning(f"Consensus fetch failed: {type(e).__name__}: {e}")

    # Generate Claude Q&A + rewrites + analysis (async, parallel)
    claude_qa = None
    claude_rewrites = None
    claude_analysis = None
    if ANTHROPIC_API_KEY:
        logger.info(f"Claude API key present, ticker={ticker}")

        # Prepare flagged sentences for rewrite — start with flagged_passages
        flagged_for_rewrite = [
            {"sentence": fp["sentence"], "issues": fp["issues"]}
            for fp in base_analysis.get("flagged_passages", [])
        ]

        # Also include sentences that classify_sentence() would flag in the Word export
        # but aren't in flagged_passages — ensures ALL highlighted sentences get rewrites
        _fp_count = len(flagged_for_rewrite)
        existing_norm = {' '.join(fp["sentence"].split()) for fp in flagged_for_rewrite}
        all_sentences = re.split(r"(?<=[.!?])\s+", transcript)
        for sent in all_sentences:
            if len(sent.strip()) < 10:
                continue
            norm_sent = ' '.join(sent.split())
            if norm_sent in existing_norm:
                continue
            color, issues = classify_sentence(sent)
            if color in ("RED", "YELLOW") and issues:
                flagged_for_rewrite.append({"sentence": sent, "issues": issues})
                existing_norm.add(norm_sent)

        logger.info(f"Sentences for rewrite: {len(flagged_for_rewrite)} total "
                     f"({_fp_count} from flagged_passages + "
                     f"{len(flagged_for_rewrite) - _fp_count} from classify_sentence)")

        # Cap at 80 sentences to manage Claude API costs/token limits
        if len(flagged_for_rewrite) > 80:
            logger.warning(f"Capping rewrite batch from {len(flagged_for_rewrite)} to 80 sentences")
            flagged_for_rewrite = flagged_for_rewrite[:80]

        # Run Claude Q&A, rewrites, analysis, and bull/bear IN PARALLEL (4 calls)
        logger.info("Calling Claude for Q&A + rewrites + analysis + bull/bear (parallel)...")

        async def _safe_qa():
            try:
                return await _generate_qa_with_claude(transcript, prior_transcripts, base_analysis)
            except Exception as e:
                logger.error(f"Claude Q&A failed: {type(e).__name__}: {e}")
                return None

        async def _safe_rewrites():
            try:
                return await _generate_rewrites_with_claude(flagged_for_rewrite)
            except Exception as e:
                logger.error(f"Claude rewrites failed: {type(e).__name__}: {e}")
                return None

        async def _safe_analysis():
            try:
                base_neg = advanced.get("negative_interpretations", [])
                base_lit = advanced.get("litigation_risk", {}).get("findings", [])
                base_act = advanced.get("activist_triggers", {}).get("triggers", [])
                return await _generate_analysis_with_claude(
                    transcript, base_neg, base_lit, base_act
                )
            except Exception as e:
                logger.error(f"Claude analysis failed: {type(e).__name__}: {e}")
                return None

        async def _safe_bull_bear():
            if not ticker or not ticker.strip():
                return None
            try:
                return await _generate_bull_bear_with_claude(transcript, ticker.strip().upper())
            except Exception as e:
                logger.error(f"Claude bull/bear failed: {type(e).__name__}: {e}")
                return None

        claude_qa, claude_rewrites, claude_analysis, claude_bull_bear = await asyncio.gather(
            _safe_qa(), _safe_rewrites(), _safe_analysis(), _safe_bull_bear()
        )

        if claude_qa:
            logger.info(
                f"Claude generated {claude_qa['total_questions']} questions "
                f"(used {claude_qa.get('prior_calls_used', 0)} prior calls)"
            )
        if claude_rewrites:
            logger.info(f"Claude generated {len(claude_rewrites)} rewrites")
        if claude_analysis:
            logger.info(f"Claude analysis: {len(claude_analysis.get('negative_interpretations', []))} neg interps, "
                         f"{len(claude_analysis.get('guidance_metrics', []))} guidance metrics")
        if claude_bull_bear:
            logger.info(f"Claude bull/bear: {len(claude_bull_bear.get('bull_cases', []))} bull, "
                         f"{len(claude_bull_bear.get('bear_cases', []))} bear, "
                         f"{len(claude_bull_bear.get('rewrites', []))} rewrites")
    else:
        logger.info("No ANTHROPIC_API_KEY — using fallback Q&A and rewrites")
        claude_bull_bear = None
        flagged_for_rewrite = []

    # Build AI status so the UI can surface degraded modes instead of silently
    # serving regex-template output when a Claude call fails.
    _has_ticker = bool(ticker and ticker.strip())
    _expected = {
        "qa": True,
        "rewrites": len(flagged_for_rewrite) > 0 if ANTHROPIC_API_KEY else False,
        "analysis": True,
        "bull_bear": _has_ticker,
    }
    _succeeded = {
        "qa": claude_qa is not None,
        "rewrites": claude_rewrites is not None,
        "analysis": claude_analysis is not None,
        "bull_bear": claude_bull_bear is not None,
    }
    _degraded_services = [
        name for name, expected in _expected.items()
        if expected and not _succeeded[name]
    ]
    ai_status = {
        "api_configured": bool(ANTHROPIC_API_KEY),
        "degraded_services": _degraded_services,
        "degraded": bool(ANTHROPIC_API_KEY) and len(_degraded_services) > 0,
    }
    if ai_status["degraded"]:
        logger.warning(f"AI degraded — falling back to templates for: {_degraded_services}")

    # Score prior transcripts for comparison (fast — no Claude calls needed)
    prior_scored = []
    if prior_transcripts:
        current_scores = _corrected_scores(base_analysis)
        for t in prior_transcripts:
            try:
                prior_base = analyze_transcript(t["content"], LM_DICT)
                prior_sc = _corrected_scores(prior_base)
                prior_scored.append({"quarter": t["quarter"], "scores": prior_sc})
            except Exception as e:
                logger.warning(f"Failed to score prior transcript {t['quarter']}: {e}")
        if prior_scored:
            logger.info(f"Scored {len(prior_scored)} prior transcripts for comparison")

    prior_comparison = _build_prior_comparison(
        _corrected_scores(base_analysis), prior_scored
    )

    # Create session
    session_id = str(uuid.uuid4())
    SESSIONS[session_id] = {
        "text": transcript,
        "base_analysis": base_analysis,
        "advanced": advanced,
        "claude_qa": claude_qa,
        "claude_rewrites": claude_rewrites,
        "claude_analysis": claude_analysis,
        "claude_bull_bear": claude_bull_bear,
        "consensus": consensus,
        "prior_comparison": prior_comparison,
        "ai_status": ai_status,
        "ticker": ticker.strip().upper() if ticker and ticker.strip() else None,
    }

    return _build_response(
        transcript, base_analysis, advanced, session_id,
        claude_qa, claude_rewrites, claude_analysis, consensus,
        prior_comparison, claude_bull_bear, ai_status,
    )


@app.get("/api/export/pdf/{session_id}")
async def get_pdf(session_id: str):
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    export_pdf(
        session["text"], session["base_analysis"], tmp.name,
        ticker=session.get("ticker"),
        prior_comparison=session.get("prior_comparison"),
    )
    return FileResponse(
        tmp.name,
        media_type="application/pdf",
        filename="streetsignals_report.pdf",
    )


@app.get("/api/export/word/{session_id}")
async def get_word(session_id: str):
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    _export_word_improved(
        session["text"],
        session["base_analysis"],
        tmp.name,
        session.get("claude_rewrites"),
        session.get("advanced"),
        session.get("claude_analysis"),
        session.get("prior_comparison"),
        session.get("claude_bull_bear"),
    )
    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="streetsignals_report.docx",
    )


@app.get("/api/export/json/{session_id}")
async def get_json(session_id: str):
    session = SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    return _build_response(
        session["text"],
        session["base_analysis"],
        session["advanced"],
        session_id,
        session.get("claude_qa"),
        session.get("claude_rewrites"),
        session.get("claude_analysis"),
        session.get("consensus"),
        session.get("prior_comparison"),
        session.get("claude_bull_bear"),
        session.get("ai_status"),
    )




# ---------------------------------------------------------------------------
# Serve React frontend (production — built files in ./static)
# ---------------------------------------------------------------------------

STATIC_DIR = Path(__file__).parent / "static"

if STATIC_DIR.is_dir():
    app.mount("/assets", StaticFiles(directory=STATIC_DIR / "assets"), name="assets")

    @app.get("/{full_path:path}")
    async def serve_spa(full_path: str):
        """Serve React app — all non-API routes return index.html."""
        index = STATIC_DIR / "index.html"
        return HTMLResponse(index.read_text())
